#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Buscador de Vacantes + Networking  —  App web local multiusuario
================================================================

Flask + SQLite, con LOGIN de usuario (guardados en users.txt con contraseña
protegida por hash). Cada usuario tiene su propio CV, perfil, búsquedas,
vacantes, puntuaciones, alertas y contactos.

Funciones:
  • Login / registro. Al registrarte subes tu CV para que el programa conozca
    tu perfil y enfoque los mensajes de networking a tu área de experiencia.
  • Escaneo de LinkedIn en segundo plano (por usuario), sin repetir vacantes.
  • Puntuación automática del CV al seleccionar cada vacante (0-100).
  • Alertas + notificación de escritorio cuando una vacante supera tu umbral.
  • Ejecutivos / RRHH / CEO de la empresa (enlaces listos).
  • CV optimizado para ATS (.docx) que conserva tu info y añade keywords.
  • Gráfico de vacantes y coincidencias por día.
  • Diccionario de habilidades editable desde Ajustes.
  • "Aplicar rápido": abre la vacante y copia el mensaje al portapapeles.

Ejecuta:  python app.py   →   http://127.0.0.1:5000

Dependencias:  pip install flask requests beautifulsoup4 python-docx plyer pypdf
"""

import os
import re
import io
import csv
import sys
import json
import math
import html
import time
import random
import hashlib
import binascii
import sqlite3
import threading
import smtplib
import urllib.parse
import webbrowser
from email.mime.text import MIMEText
from functools import wraps
from datetime import datetime, timedelta

import requests
from flask import (Flask, request, jsonify, send_file, Response,
                   render_template, session)

# ============================================================================
#  RUTAS Y CONFIG
# ============================================================================

APP_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(APP_DIR, "vacantes.db")
USERS_TXT = os.path.join(APP_DIR, "users.txt")
SECRET_FILE = os.path.join(APP_DIR, ".secret")
GEN_DIR = os.path.join(APP_DIR, "cv_generados")
ORIG_DIR = os.path.join(APP_DIR, "cv_originales")
AVATAR_DIR = os.path.join(APP_DIR, "avatars")
STATIC_DIR = os.path.join(APP_DIR, "static")
os.makedirs(GEN_DIR, exist_ok=True)
os.makedirs(ORIG_DIR, exist_ok=True)
os.makedirs(AVATAR_DIR, exist_ok=True)
os.makedirs(STATIC_DIR, exist_ok=True)

DEFAULTS = {
    "profile_nombre": "",
    "profile_titulo": "Profesional",
    "profile_desc_corto": "mi área de experiencia",
    "profile_desc_largo": "mi experiencia profesional",
    "cv_text": "",
    "cv_filename": "",
    "cv_ext": "",
    "scan_interval_min": "30",
    "scan_enabled": "0",
    "match_threshold": "60",
    "max_per_search": "25",
    "skill_dict": "",   # JSON del usuario; vacío = usa el diccionario base
    # Notificaciones por correo (SMTP)
    "email_enabled": "0",
    "email_to": "",
    "smtp_host": "smtp.gmail.com",
    "smtp_port": "587",
    "smtp_user": "",
    "smtp_pass": "",
    # Telegram
    "telegram_enabled": "0",
    "telegram_token": "",
    "telegram_chat_id": "",
    # Portales de empleo adicionales (además de LinkedIn)
    "sources": "remoteok,remotive,arbeitnow",
    "adzuna_app_id": "",
    "adzuna_app_key": "",
    "adzuna_country": "es",
    "rss_feeds": "",           # URLs de feeds RSS/Atom, uno por línea
    # Recordatorios de seguimiento
    "followup_days": "5",
    # Resumen (digest) por correo/Telegram
    "digest_enabled": "0",
    "digest_channel": "email",   # email | telegram | both
    "digest_freq": "daily",      # daily | weekly
    "last_digest": "",
    # Temas personalizados guardados en la cuenta (JSON: lista de objetos)
    "custom_themes": "",
    # Meta semanal de postulaciones (0 = desactivada)
    "weekly_goal": "10",
    # Redacción de mensajes con IA
    "ai_enabled": "0",
    "ai_provider": "ollama",     # ollama | openai | anthropic | gemini | custom
    "ai_base_url": "",           # para ollama/custom (ej. http://localhost:11434)
    "ai_model": "",              # ej. llama3.1, gpt-4o-mini, claude-3-5-haiku, gemini-1.5-flash
    "ai_api_key": "",            # secreto (no se envía al cliente)
    # Bloc de notas + pizarra (guardados en la cuenta)
    "scratchpad": "",
    "board_png": "",
}

# Portales disponibles (clave -> nombre visible)
SOURCE_NAMES = {
    "linkedin": "LinkedIn",
    "remoteok": "RemoteOK",
    "remotive": "Remotive",
    "arbeitnow": "Arbeitnow",
    "adzuna": "Adzuna",
    "rss": "RSS",
}

HTTP_HEADERS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
                   "(KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36"),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "es-ES,es;q=0.9,en;q=0.8",
}

# ---- Diccionario base de habilidades (soporte + ciberseguridad + dev) ----
BASE_SKILLS = {
    "javascript": ["javascript", "js"], "typescript": ["typescript", "ts"],
    "react": ["react", "reactjs"], "node.js": ["node.js", "nodejs", "node"],
    "sql": ["sql", "postgres", "postgresql", "mysql", "sqlite"],
    "python": ["python"], "html": ["html", "html5"],
    "css": ["css", "tailwind", "sass", "bootstrap"],
    "git": ["git", "github", "gitlab", "control de versiones"],
    "rest api": ["rest", "api rest", "restful", "apis"],
    "docker": ["docker", "contenedor", "container"],
    "agile": ["agile", "scrum", "kanban", "metodologias agiles"],
    "ci/cd": ["ci/cd", "cicd", "integracion continua", "jenkins"],
    "liderazgo": ["liderazgo", "lead", "leadership", "gestion de equipos"],
    "ingles": ["ingles", "english", "b2", "c1", "bilingue"],
    "soporte tecnico": ["soporte tecnico", "soporte", "help desk", "helpdesk",
                          "mesa de ayuda", "service desk", "technical support",
                          "soporte a usuarios", "soporte n1", "soporte n2"],
    "troubleshooting": ["troubleshooting", "resolucion de problemas",
                         "resolucion de incidencias", "diagnostico"],
    "itil": ["itil", "gestion de servicios", "sla"],
    "ticketing": ["ticketing", "tickets", "zendesk", "freshdesk", "glpi",
                  "osticket", "servicenow"],
    "redes": ["redes", "networking", "tcp/ip", "lan", "wan", "vpn", "dns",
              "dhcp", "router", "switch", "vlan"],
    "windows": ["windows", "windows server"],
    "linux": ["linux", "ubuntu", "debian", "centos", "red hat", "bash"],
    "active directory": ["active directory", "directorio activo", "ldap",
                         "gpo", "azure ad", "entra id"],
    "hardware": ["hardware", "ensamblaje", "mantenimiento de equipos",
                 "perifericos", "impresoras"],
    "ofimatica": ["office 365", "microsoft 365", "outlook", "ofimatica",
                  "google workspace", "sharepoint"],
    "ciberseguridad": ["ciberseguridad", "seguridad informatica",
                        "seguridad de la informacion", "cybersecurity",
                        "infosec", "seguridad", "security"],
    "iso 27001": ["iso 27001", "iso27001", "sgsi", "isms", "anexo a"],
    "gestion de riesgos": ["gestion de riesgos", "risk management",
                           "analisis de riesgos", "evaluacion de riesgos"],
    "siem": ["siem", "splunk", "wazuh", "qradar", "graylog"],
    "firewall": ["firewall", "fortinet", "palo alto", "pfsense", "cortafuegos",
                 "ips", "ids", "fortigate"],
    "edr / antivirus": ["edr", "antivirus", "endpoint protection", "crowdstrike",
                        "defender", "sentinelone", "malware", "ransomware", "xdr"],
    "pentesting": ["pentesting", "ethical hacking", "hacking etico", "kali",
                   "nmap", "metasploit", "burp", "owasp"],
    "vulnerabilidades": ["vulnerabilidades", "nessus", "openvas", "cve",
                         "escaneo de vulnerabilidades", "gestion de parches"],
    "hardening": ["hardening", "bastionado", "endurecimiento"],
    "respuesta a incidentes": ["respuesta a incidentes", "incident response",
                               "soc", "analista soc", "forense", "dfir"],
    "cumplimiento / normativas": ["cumplimiento", "compliance", "gdpr", "rgpd",
                                  "nist", "ens", "pci dss", "hipaa", "auditoria",
                                  "auditor"],
    "backup / continuidad": ["backup", "copias de seguridad", "respaldo",
                             "recuperacion", "disaster recovery", "continuidad"],
    "cloud": ["azure", "aws", "gcp", "google cloud", "intune", "cloud"],
    "certificaciones sec": ["comptia", "security+", "network+", "ceh", "ccna",
                            "itil foundation", "cissp"],
}

STOPWORDS = set("""
a al algo algun alguna alguno algunos ante antes como con contra cual cuando de
del desde donde dos el ella ellos en entre era es esa ese eso esta estar este esto
estos fin fue fueron ha hace hacia han hasta hay la las le les lo los mas me mi mis
mucho muy nada ni no nos nuestra nuestro o os otra otro para pero poco por porque
que quien se sea segun ser si sin sobre solo son su sus tal tambien tanto te tiene
toda todo todos tu tus un una uno unos usted y ya the a an and or but of to in on
for with at by from up about into over after is are was were be been being this that
these those it its as your you we our their they he she his her will can would job
role work team company
""".split())


# ============================================================================
#  TEXTO / SCORING
# ============================================================================

def tokenize(text):
    tokens = re.findall(r"[a-záéíóúüñ0-9\+#\.]{2,}", text.lower())
    return [t for t in tokens if t not in STOPWORDS and len(t) > 1]


def tfidf_cosine(a, b):
    ta, tb = tokenize(a), tokenize(b)
    if not ta or not tb:
        return 0.0

    def tf(tokens):
        d = {}
        for t in tokens:
            d[t] = d.get(t, 0) + 1
        n = len(tokens)
        return {t: c / n for t, c in d.items()}

    fa, fb = tf(ta), tf(tb)
    vocab = set(fa) | set(fb)
    idf = {t: math.log(3 / ((1 if t in fa else 0) + (1 if t in fb else 0) + 1)) + 1
           for t in vocab}
    va = {t: fa.get(t, 0) * idf[t] for t in vocab}
    vb = {t: fb.get(t, 0) * idf[t] for t in vocab}
    dot = sum(va[t] * vb[t] for t in vocab)
    ma = math.sqrt(sum(v * v for v in va.values()))
    mb = math.sqrt(sum(v * v for v in vb.values()))
    return dot / (ma * mb) if ma and mb else 0.0


def detect_skills(text, skill_dict=None):
    sd = skill_dict or BASE_SKILLS
    low = " " + text.lower() + " "
    found = set()
    for canon, variants in sd.items():
        for v in variants:
            if re.search(r"(?<![a-z0-9])" + re.escape(v) + r"(?![a-z0-9])", low):
                found.add(canon)
                break
    return found


def score_cv(cv_text, job_text, skill_dict=None):
    if not cv_text.strip():
        return None
    sim = tfidf_cosine(cv_text, job_text)
    js = detect_skills(job_text, skill_dict)
    cs = detect_skills(cv_text, skill_dict)
    matched = sorted(js & cs)
    missing = sorted(js - cs)
    extra = sorted(cs - js)
    ratio = len(matched) / len(js) if js else 0.0
    simc = min(1.0, sim * 2.2)
    score = round(min(100.0, (simc * 0.70 + ratio * 0.30) * 100), 1)
    if score >= 75:
        verdict = "Excelente encaje"
    elif score >= 55:
        verdict = "Buen encaje"
    elif score >= 35:
        verdict = "Encaje parcial"
    else:
        verdict = "Encaje bajo"
    return {"score": score, "similarity_pct": round(simc * 100, 1),
            "skill_pct": round(ratio * 100, 1), "matched": matched,
            "missing": missing, "extra": extra, "verdict": verdict}


def parse_skill_dict(raw):
    """Convierte el JSON del usuario a dict; si falla, usa el base."""
    if not raw or not raw.strip():
        return dict(BASE_SKILLS)
    try:
        d = json.loads(raw)
        if isinstance(d, dict) and d:
            return {str(k): [str(x) for x in v] for k, v in d.items()}
    except Exception:
        pass
    return dict(BASE_SKILLS)


# ============================================================================
#  LECTURA DE CV
# ============================================================================

def extract_cv_text(filename, data_bytes):
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".txt", ".md"):
        return data_bytes.decode("utf-8", errors="ignore")
    if ext == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("Instala 'pypdf' para leer PDF (pip install pypdf).")
        reader = PdfReader(io.BytesIO(data_bytes))
        text = "\n".join((p.extract_text() or "") for p in reader.pages).strip()
        if not text:
            raise RuntimeError("PDF sin texto legible (¿escaneado?). Usa .docx o .txt.")
        return text
    if ext == ".docx":
        import docx
        doc = docx.Document(io.BytesIO(data_bytes))
        return "\n".join(p.text for p in doc.paragraphs)
    raise RuntimeError(f"Formato no soportado: {ext}. Usa PDF, DOCX o TXT.")


# ============================================================================
#  CLIENTE LINKEDIN
# ============================================================================

class LinkedIn:
    SEARCH = "https://www.linkedin.com/jobs-guest/jobs/api/seeMoreJobPostings/search"
    JOB = "https://www.linkedin.com/jobs-guest/jobs/api/jobPosting/{jid}"

    def __init__(self):
        self.s = requests.Session()
        self.s.headers.update(HTTP_HEADERS)

    def search(self, keyword, location, max_results=25):
        results, seen, start = [], set(), 0
        while len(results) < max_results and start < max_results + 20:
            params = {"keywords": keyword, "location": location, "start": start,
                      "sortBy": "DD", "f_TPR": "r2592000"}
            try:
                r = self.s.get(self.SEARCH, params=params, timeout=20)
            except Exception:
                break
            if r.status_code != 200:
                break
            cards = self._cards(r.text)
            if not cards:
                break
            added = 0
            for c in cards:
                if c["id"] in seen:
                    continue
                seen.add(c["id"])
                results.append(c)
                added += 1
                if len(results) >= max_results:
                    break
            if not added:
                break
            start += 10
            time.sleep(random.uniform(0.8, 1.5))
        return results[:max_results]

    def _cards(self, text):
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return self._cards_regex(text)
        soup = BeautifulSoup(text, "html.parser")
        out = []
        for li in soup.find_all("li"):
            t = li.find(class_="base-search-card__title")
            comp = li.find(class_="base-search-card__subtitle")
            loc = li.find(class_="job-search-card__location")
            link = li.find("a", class_="base-card__full-link") or li.find("a")
            tm = li.find("time")
            base = li.find(class_="base-card")
            jid = None
            if base and base.get("data-entity-urn"):
                jid = base["data-entity-urn"].split(":")[-1]
            url = link.get("href").split("?")[0] if link and link.get("href") else ""
            if not jid and url:
                m = re.search(r"-(\d+)$", url)
                jid = m.group(1) if m else url
            if not (t and jid):
                continue
            out.append({"id": str(jid), "title": t.get_text(strip=True),
                        "company": comp.get_text(strip=True) if comp else "—",
                        "location": loc.get_text(strip=True) if loc else "",
                        "url": url,
                        "date": (tm.get("datetime") if tm and tm.get("datetime") else "")})
        return out

    def _cards_regex(self, text):
        out = []
        for b in re.split(r'<li[ >]', text)[1:]:
            idm = re.search(r'urn:li:jobPosting:(\d+)', b)
            tm = re.search(r'base-search-card__title[^>]*>\s*(.*?)\s*<', b, re.S)
            cm = re.search(r'base-search-card__subtitle[^>]*>\s*(?:<a[^>]*>)?\s*(.*?)\s*<', b, re.S)
            lm = re.search(r'job-search-card__location[^>]*>\s*(.*?)\s*<', b, re.S)
            um = re.search(r'href="(https://[^"]*?/jobs/view/[^"?]+)', b)
            dm = re.search(r'datetime="([^"]+)"', b)
            if not (idm and tm):
                continue
            out.append({"id": idm.group(1), "title": html.unescape(tm.group(1).strip()),
                        "company": html.unescape(cm.group(1).strip()) if cm else "—",
                        "location": html.unescape(lm.group(1).strip()) if lm else "",
                        "url": um.group(1) if um else "",
                        "date": dm.group(1) if dm else ""})
        return out

    def details(self, jid):
        try:
            r = self.s.get(self.JOB.format(jid=jid), timeout=20)
        except Exception as e:
            return f"(No se pudo cargar: {e})", None
        if r.status_code != 200:
            return f"(No se pudo cargar: HTTP {r.status_code})", None
        text = r.text
        easy = self._easy(text)
        crit = self._criteria(text)
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(text, "html.parser")
            el = (soup.find(class_="show-more-less-html__markup")
                  or soup.find(class_="description__text"))
            body = el.get_text("\n", strip=True) if el else ""
        except ImportError:
            m = re.search(r'show-more-less-html__markup[^>]*>(.*?)</div>', text, re.S)
            body = html.unescape(re.sub(r"<[^>]+>", " ", m.group(1))).strip() if m else ""
        if not body:
            body = "(LinkedIn no devolvió la descripción para invitados.)"
        return ((crit + "\n\n" + body).strip() if crit else body), easy

    @staticmethod
    def _easy(text):
        low = text.lower()
        if 'id="applyurl"' in low or "apply-button--link" in low:
            return 0
        if ("easy apply" in low or "solicitud sencilla" in low
                or "postulación sencilla" in low or "postulacion sencilla" in low):
            return 1
        if "apply-button" in low:
            return 1
        return None

    @staticmethod
    def _criteria(text):
        pairs = re.findall(
            r'description__job-criteria-subheader[^>]*>\s*(.*?)\s*<.*?'
            r'description__job-criteria-text[^>]*>\s*(.*?)\s*<', text, re.S)
        lines = []
        for k, v in pairs:
            k = html.unescape(re.sub(r"<[^>]+>", "", k)).strip()
            v = html.unescape(re.sub(r"<[^>]+>", "", v)).strip()
            if k and v:
                lines.append(f"• {k}: {v}")
        return "\n".join(lines)


LI = LinkedIn()


# ============================================================================
#  OTROS PORTALES DE EMPLEO (APIs públicas)
# ============================================================================

def _strip_html(t):
    return html.unescape(re.sub(r"<[^>]+>", " ", t or "")).strip()


def _kw_match(keyword, *texts):
    if not keyword:
        return True
    blob = " ".join(t or "" for t in texts).lower()
    return all(w in blob for w in keyword.lower().split())


def _norm_job(source, oid, title, company, location, url, date, desc):
    return {"id": f"{source}:{oid}", "source": source, "title": title or "—",
            "company": company or "—", "location": location or "Remoto",
            "url": url or "", "date": (date or "")[:10], "description": desc or "",
            "easy_apply": None}


def src_remoteok(keyword, location, limit):
    r = requests.get("https://remoteok.com/api", headers=HTTP_HEADERS, timeout=20)
    out = []
    for it in r.json():
        if not isinstance(it, dict) or "id" not in it:
            continue
        title = it.get("position") or it.get("title", "")
        tags = " ".join(it.get("tags", []) or [])
        if not _kw_match(keyword, title, tags, it.get("description", "")):
            continue
        out.append(_norm_job("remoteok", it["id"], title, it.get("company", ""),
                             it.get("location") or "Remoto",
                             it.get("url") or it.get("apply_url"),
                             it.get("date", ""), _strip_html(it.get("description", ""))))
        if len(out) >= limit:
            break
    return out


def src_remotive(keyword, location, limit):
    r = requests.get("https://remotive.com/api/remote-jobs",
                     params={"search": keyword, "limit": limit},
                     headers=HTTP_HEADERS, timeout=20)
    out = []
    for it in r.json().get("jobs", []):
        out.append(_norm_job("remotive", it["id"], it.get("title", ""),
                             it.get("company_name", ""),
                             it.get("candidate_required_location") or "Remoto",
                             it.get("url"), it.get("publication_date", ""),
                             _strip_html(it.get("description", ""))))
        if len(out) >= limit:
            break
    return out


def src_arbeitnow(keyword, location, limit):
    r = requests.get("https://www.arbeitnow.com/api/job-board-api",
                     headers=HTTP_HEADERS, timeout=20)
    out = []
    for it in r.json().get("data", []):
        tags = " ".join(it.get("tags", []) or [])
        if not _kw_match(keyword, it.get("title", ""), it.get("description", ""), tags):
            continue
        out.append(_norm_job("arbeitnow", it.get("slug", it.get("url", "")),
                             it.get("title", ""), it.get("company_name", ""),
                             it.get("location") or ("Remoto" if it.get("remote") else ""),
                             it.get("url"), it.get("created_at", ""),
                             _strip_html(it.get("description", ""))))
        if len(out) >= limit:
            break
    return out


def src_adzuna(keyword, location, limit, app_id="", app_key="", country="es"):
    if not (app_id and app_key):
        return []
    url = f"https://api.adzuna.com/v1/api/jobs/{country}/search/1"
    r = requests.get(url, params={"app_id": app_id, "app_key": app_key,
                                  "what": keyword, "where": location or "",
                                  "results_per_page": limit},
                     headers=HTTP_HEADERS, timeout=20)
    out = []
    for it in r.json().get("results", []):
        out.append(_norm_job("adzuna", it.get("id"), it.get("title", ""),
                             (it.get("company") or {}).get("display_name", ""),
                             (it.get("location") or {}).get("display_name", ""),
                             it.get("redirect_url"), it.get("created", ""),
                             _strip_html(it.get("description", ""))))
        if len(out) >= limit:
            break
    return out


def src_rss(feeds, keyword, limit):
    """Lee feeds RSS/Atom genéricos (pega la URL de cualquier bolsa de empleo)."""
    import xml.etree.ElementTree as ET
    out = []
    for url in feeds:
        url = url.strip()
        if not url:
            continue
        try:
            r = requests.get(url, headers=HTTP_HEADERS, timeout=20)
            root = ET.fromstring(r.content)
        except Exception:
            continue
        # RSS <item> o Atom <entry>
        items = root.iter("item")
        items = list(items) or [e for e in root.iter() if e.tag.endswith("entry")]
        host = urllib.parse.urlparse(url).netloc.replace("www.", "")
        for it in items:
            def g(tag):
                for ch in it:
                    if ch.tag.endswith(tag):
                        if tag == "link" and not (ch.text or "").strip():
                            return ch.get("href", "")
                        return (ch.text or "").strip()
                return ""
            title = g("title")
            desc = _strip_html(g("description") or g("summary") or g("content"))
            link = g("link")
            date = g("pubDate") or g("updated") or g("published")
            if not title:
                continue
            if not _kw_match(keyword, title, desc):
                continue
            oid = g("guid") or link or title
            out.append(_norm_job("rss", oid[:120], title, host, "", link, date, desc))
            if len(out) >= limit:
                break
    return out


def fetch_sources(enabled, keyword, location, limit, adzuna=None, rss=None, log=None):
    """Consulta los portales habilitados y devuelve vacantes normalizadas."""
    jobs = []
    fns = {"remoteok": src_remoteok, "remotive": src_remotive, "arbeitnow": src_arbeitnow}
    for key in enabled:
        try:
            if key == "adzuna" and adzuna:
                jobs += src_adzuna(keyword, location, limit, **adzuna)
            elif key == "rss" and rss:
                jobs += src_rss(rss, keyword, limit)
            elif key in fns:
                jobs += fns[key](keyword, location, limit)
        except Exception as e:
            if log:
                log(f"{SOURCE_NAMES.get(key, key)}: {type(e).__name__}")
    return jobs


# ============================================================================
#  TELEGRAM
# ============================================================================

def send_telegram(token, chat_id, text):
    if not (token and chat_id):
        raise RuntimeError("Falta el token o el chat_id de Telegram.")
    r = requests.post(f"https://api.telegram.org/bot{token}/sendMessage",
                      data={"chat_id": chat_id, "text": text,
                            "disable_web_page_preview": False}, timeout=15)
    if r.status_code != 200:
        raise RuntimeError(f"Telegram respondió {r.status_code}: {r.text[:120]}")


# ============================================================================
#  DECISORES Y MENSAJES (enfocados al área del usuario)
# ============================================================================

def people_links(company):
    q = urllib.parse.quote(company)
    roles = urllib.parse.quote(f'{company} (CEO OR Fundador OR Director OR Gerente OR CTO)')
    rrhh = urllib.parse.quote(f'{company} (Recursos Humanos OR Talent OR Reclutador OR Recruiter)')
    xray = urllib.parse.quote(f'site:linkedin.com/in "{company}" (CEO OR Director OR '
                              f'"Recursos Humanos" OR Talent OR Fundador OR Gerente)')
    return [
        {"label": f"Ejecutivos / altos cargos de «{company}»",
         "url": f"https://www.linkedin.com/search/results/people/?keywords={roles}"},
        {"label": f"RRHH / Reclutamiento de «{company}»",
         "url": f"https://www.linkedin.com/search/results/people/?keywords={rrhh}"},
        {"label": f"Todos los empleados de «{company}»",
         "url": f"https://www.linkedin.com/search/results/people/?keywords={q}"},
        {"label": f"Google X-ray de decisores de «{company}»",
         "url": f"https://www.google.com/search?q={xray}"},
        {"label": f"Página de «{company}» en LinkedIn",
         "url": f"https://www.linkedin.com/search/results/companies/?keywords={q}"},
    ]


def recruiter_messages(profile, company, role, matched, top_skills, contact="",
                       lang="es", keywords=None):
    """Genera mensajes al reclutador en español o inglés, incorporando las
    palabras clave que el usuario quiera resaltar. Se regeneran en cada llamada."""
    nombre = profile.get("profile_nombre") or ("yo" if lang == "es" else "me")
    titulo = profile.get("profile_titulo") or ("profesional" if lang == "es" else "professional")
    corto = profile.get("profile_desc_corto") or ("mi área" if lang == "es" else "my field")
    largo = profile.get("profile_desc_largo") or ("mi experiencia" if lang == "es" else "my experience")
    contact = (contact or "").strip()
    first = contact.split()[0] if contact else ""
    # Palabras clave que pidió el usuario (prioridad); si no, las que coinciden; si no, del CV
    kws = [k.strip() for k in (keywords or []) if k and k.strip()]
    focus = kws or matched[:5] or (top_skills[:5] if top_skills else [])

    if lang == "en":
        role = role or "the role"
        company = company or "your company"
        hola = f"Hi {first}" if first else "Hi"
        holae = f"Hi {first}!" if first else "Hi!"
        formal = f"Dear {contact}" if contact else f"Dear {company} team"
        # En inglés NO incrustamos los descriptores del perfil (que suelen estar en
        # español); nos apoyamos en el puesto y en las skills (términos técnicos).
        strengths = ("hands-on experience with " + ", ".join(focus)) if focus else "a solid background in this area"
        kw_line = (" I'd especially highlight: " + ", ".join(kws) + ".") if kws else ""
        return [
            {"title": "Connection note (short)",
             "body": f"{hola}, I came across the {role} opening at {company} and it's a great fit for my profile. "
                     f"I'd love to connect and chat about how I can add value.{kw_line} Thanks!"},
            {"title": "Formal message",
             "body": f"{formal},\n\nI'm reaching out about the {role} position — it really caught my "
                     f"interest. I bring {strengths}, and I believe my profile aligns "
                     f"well with what you're looking for.{kw_line}\n\nI'd love to talk and share my CV. "
                     f"Looking forward to hearing from you.\n\nBest regards,\n{nombre}"},
            {"title": "Warm, value-first",
             "body": f"{holae} I saw {company} is hiring a {role} and it really stood out to me. "
                     f"I've been working with {strengths}, and I'd be excited to join a team like yours.{kw_line}"
                     f"\n\nWould you be open to a short chat? Happy to send my CV. Thanks for your time!"},
            {"title": "Exploratory networking",
             "body": f"{hola}, I follow {company}'s work and would love to connect. "
                     f"I'm exploring new opportunities and bring {strengths}.{kw_line} If you're looking for someone "
                     f"with my profile, I'd love to be considered. Cheers!"},
        ]

    # Español (por defecto)
    role = role or "la vacante publicada"
    company = company or "su empresa"
    hola = f"Hola {first}" if first else "Hola"
    holae = f"¡Hola {first}!" if first else "¡Hola!"
    formal = f"Estimado/a {contact}" if contact else f"Estimado/a equipo de {company}"
    strengths = ("experiencia en " + ", ".join(focus)) if focus else largo
    kw_line = (" Destacaría especialmente: " + ", ".join(kws) + ".") if kws else ""
    return [
        {"title": "Nota de conexión (breve)",
         "body": f"{hola}, vi la vacante de {role} en {company} y me encajó mucho. "
                 f"Trabajo en {corto} y me encantaría conectar para conversar sobre "
                 f"cómo puedo aportar.{kw_line} ¡Gracias!"},
        {"title": "Mensaje formal",
         "body": f"{formal}:\n\nMe pongo en contacto porque vi la publicación de "
                 f"{role} y me interesa mucho la oportunidad. Me desempeño como "
                 f"{titulo}, con {strengths}, y creo que mi perfil se alinea bien con "
                 f"lo que buscan.{kw_line}\n\nMe encantaría conversar y compartirles mi CV. "
                 f"Quedo atento/a a su respuesta.\n\nSaludos cordiales,\n{nombre}"},
        {"title": "Cercano con propuesta de valor",
         "body": f"{holae} Vi que en {company} buscan {role} y me llamó mucho la "
                 f"atención. Justo vengo trabajando con {strengths}, y me entusiasma "
                 f"sumar a un equipo como el suyo.{kw_line}\n\n¿Estarían abiertos a una breve "
                 f"conversación? Con gusto les paso mi CV. ¡Gracias por su tiempo!"},
        {"title": "Networking exploratorio",
         "body": f"{hola}, sigo el trabajo de {company} y me gustaría conectar. Me "
                 f"desempeño en {corto} y estoy explorando nuevas oportunidades.{kw_line} Si "
                 f"buscan a alguien con mi perfil, me encantaría que me tuvieran en "
                 f"cuenta. ¡Un saludo!"},
    ]


# ============================================================================
#  REDACCIÓN CON IA (opcional): Ollama local / OpenAI / Anthropic / Gemini
# ============================================================================
AI_TIMEOUT = 60


def ai_chat(cfg, system, user, max_tokens=700):
    """Envía system+user al proveedor configurado y devuelve el texto.
    Lanza Exception con un mensaje claro si algo falla."""
    provider = (cfg.get("ai_provider") or "ollama").strip().lower()
    model = (cfg.get("ai_model") or "").strip()
    key = (cfg.get("ai_api_key") or "").strip()
    base = (cfg.get("ai_base_url") or "").strip().rstrip("/")

    if provider == "ollama":
        url = (base or "http://localhost:11434") + "/api/chat"
        payload = {"model": model or "llama3.1", "stream": False,
                   "messages": [{"role": "system", "content": system},
                                {"role": "user", "content": user}]}
        r = requests.post(url, json=payload, timeout=AI_TIMEOUT)
        r.raise_for_status()
        return (r.json().get("message", {}) or {}).get("content", "").strip()

    if provider in ("openai", "custom"):
        if not key and provider == "openai":
            raise Exception("Falta la API key de OpenAI.")
        url = (base or "https://api.openai.com/v1").rstrip("/") + "/chat/completions"
        headers = {"Authorization": f"Bearer {key}", "Content-Type": "application/json"}
        payload = {"model": model or "gpt-4o-mini", "max_tokens": max_tokens,
                   "messages": [{"role": "system", "content": system},
                                {"role": "user", "content": user}]}
        r = requests.post(url, json=payload, headers=headers, timeout=AI_TIMEOUT)
        r.raise_for_status()
        return r.json()["choices"][0]["message"]["content"].strip()

    if provider == "anthropic":
        if not key:
            raise Exception("Falta la API key de Anthropic.")
        url = (base or "https://api.anthropic.com") + "/v1/messages"
        headers = {"x-api-key": key, "anthropic-version": "2023-06-01",
                   "Content-Type": "application/json"}
        payload = {"model": model or "claude-3-5-haiku-latest", "max_tokens": max_tokens,
                   "system": system, "messages": [{"role": "user", "content": user}]}
        r = requests.post(url, json=payload, headers=headers, timeout=AI_TIMEOUT)
        r.raise_for_status()
        return "".join(b.get("text", "") for b in r.json().get("content", [])).strip()

    if provider == "gemini":
        if not key:
            raise Exception("Falta la API key de Gemini.")
        m = model or "gemini-1.5-flash"
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{m}:generateContent?key={key}"
        payload = {"system_instruction": {"parts": [{"text": system}]},
                   "contents": [{"parts": [{"text": user}]}],
                   "generationConfig": {"maxOutputTokens": max_tokens}}
        r = requests.post(url, json=payload, timeout=AI_TIMEOUT)
        r.raise_for_status()
        cand = r.json().get("candidates", [{}])[0]
        return "".join(p.get("text", "") for p in cand.get("content", {}).get("parts", [])).strip()

    raise Exception(f"Proveedor de IA desconocido: {provider}")


def ai_message_prompt(profile, job, lang, keywords, tone, contact=""):
    """Construye el system+user para que la IA redacte el mensaje al reclutador."""
    nombre = profile.get("profile_nombre") or ""
    titulo = profile.get("profile_titulo") or ""
    corto = profile.get("profile_desc_corto") or ""
    largo = profile.get("profile_desc_largo") or ""
    desc = (job.get("description") or "")[:1500]
    kws = ", ".join(keywords) if keywords else ""
    tones = {
        "breve": {"es": "una nota de conexión breve (2-3 frases, para solicitud de conexión en LinkedIn)",
                  "en": "a short connection note (2-3 sentences, for a LinkedIn connection request)"},
        "formal": {"es": "un mensaje formal y profesional (un párrafo o dos)",
                   "en": "a formal, professional message (one or two paragraphs)"},
        "cercano": {"es": "un mensaje cercano y con energía, destacando tu propuesta de valor",
                    "en": "a warm, energetic message highlighting your value proposition"},
    }
    style = tones.get(tone, tones["formal"])[("en" if lang == "en" else "es")]

    if lang == "en":
        system = ("You are an expert career coach who writes concise, natural, and persuasive "
                  "outreach messages that job seekers send to recruiters on LinkedIn. Write ONLY the "
                  "message body, ready to send. Do not include a subject line, placeholders in brackets, "
                  "or any explanation. Never invent experience the candidate did not provide. Write in fluent English.")
        user = (f"Write {style} to a recruiter about this job.\n\n"
                f"JOB TITLE: {job.get('title','')}\nCOMPANY: {job.get('company','')}\n"
                f"JOB DESCRIPTION: {desc}\n\n"
                f"CANDIDATE NAME: {nombre}\nCANDIDATE ROLE: {titulo}\n"
                f"CANDIDATE BACKGROUND: {corto}. {largo}\n"
                f"SKILLS THAT MATCH: {job.get('matched','')}\n")
        if kws:
            user += f"MUST NATURALLY HIGHLIGHT THESE: {kws}\n"
        if contact:
            user += f"RECRUITER'S NAME (address them by first name): {contact}\n"
        user += "\nSign it with the candidate's name if appropriate. Keep it human, specific to the company, and not generic."
        return system, user

    system = ("Eres un coach de carrera experto que redacta mensajes de contacto breves, naturales y "
              "persuasivos que un candidato envía a un reclutador por LinkedIn. Escribe SOLO el cuerpo del "
              "mensaje, listo para enviar. No incluyas asunto, ni marcadores entre corchetes, ni explicaciones. "
              "Nunca inventes experiencia que el candidato no haya indicado. Escribe en español natural.")
    user = (f"Redacta {style} para un reclutador sobre esta vacante.\n\n"
            f"VACANTE: {job.get('title','')}\nEMPRESA: {job.get('company','')}\n"
            f"DESCRIPCIÓN: {desc}\n\n"
            f"NOMBRE DEL CANDIDATO: {nombre}\nROL DEL CANDIDATO: {titulo}\n"
            f"PERFIL: {corto}. {largo}\n"
            f"SKILLS QUE COINCIDEN: {job.get('matched','')}\n")
    if kws:
        user += f"RESALTA ESTO DE FORMA NATURAL: {kws}\n"
    if contact:
        user += f"NOMBRE DEL RECLUTADOR (salúdalo por su nombre): {contact}\n"
    user += "\nFírmalo con el nombre del candidato si procede. Que sea humano, específico a la empresa y nada genérico."
    return system, user


def ai_score_cv(cfg, cv_text, job_title, job_desc):
    """Puntúa el CV contra la vacante con IA. Devuelve dict como score_cv o None si falla."""
    system = ("Eres un reclutador técnico. Evalúas el encaje de un CV con una vacante. "
              "Responde SOLO con un JSON válido, sin texto extra, con esta forma exacta: "
              '{"score": <0-100 entero>, "verdict": "<3-6 palabras>", '
              '"matched": ["skill",...], "missing": ["skill",...]}')
    user = (f"VACANTE: {job_title}\nDESCRIPCIÓN:\n{(job_desc or '')[:2500]}\n\n"
            f"CV DEL CANDIDATO:\n{(cv_text or '')[:4000]}\n\n"
            "Puntúa de 0 a 100 el encaje real (experiencia + skills). 'matched' = requisitos "
            "de la vacante que el CV SÍ cumple; 'missing' = los que faltan. Máx 10 en cada lista.")
    txt = ai_chat(cfg, system, user, max_tokens=500)
    m = re.search(r"\{.*\}", txt, re.S)
    data = json.loads(m.group(0) if m else txt)
    sc = int(round(float(data.get("score", 0))))
    sc = max(0, min(100, sc))
    return {"score": sc, "verdict": str(data.get("verdict", "")).strip() or "Evaluado por IA",
            "matched": [str(x).strip() for x in (data.get("matched") or []) if str(x).strip()][:12],
            "missing": [str(x).strip() for x in (data.get("missing") or []) if str(x).strip()][:12],
            "ai": True}


def cover_letter(profile, company, role, matched, top_skills, city="", desc=""):
    """Genera una carta de presentación (cover letter) enfocada a la vacante."""
    nombre = profile.get("profile_nombre") or "[Tu nombre]"
    titulo = profile.get("profile_titulo") or "profesional"
    corto = profile.get("profile_desc_corto") or "mi área"
    largo = profile.get("profile_desc_largo") or "mi experiencia profesional"
    company = company or "su empresa"
    role = role or "la vacante"
    fuertes = matched[:6] if matched else (top_skills[:6] if top_skills else [])
    fuertes_txt = (", ".join(fuertes)) if fuertes else corto
    hoy = datetime.now().strftime("%d/%m/%Y")
    return (
        f"{hoy}\n\n"
        f"Estimado equipo de {company}:\n\n"
        f"Me dirijo a ustedes con gran interés en la posición de {role}. Como "
        f"{titulo}, con {largo}, considero que mi perfil encaja con lo que buscan.\n\n"
        f"A lo largo de mi trayectoria he trabajado con {fuertes_txt}, lo que me ha "
        f"permitido resolver problemas reales y aportar valor de forma concreta. Me "
        f"motiva especialmente la oportunidad de sumar a {company} y contribuir desde "
        f"el primer día con {corto}.\n\n"
        f"Me encantaría conversar sobre cómo puedo ayudar a su equipo a alcanzar sus "
        f"objetivos. Adjunto mi CV y quedo a su entera disposición para una "
        f"entrevista.\n\n"
        f"Agradezco su tiempo y consideración.\n\n"
        f"Atentamente,\n{nombre}")


# Preguntas técnicas por habilidad (para la preparación de entrevista)
SKILL_QUESTIONS = {
    "iso 27001": "¿Cómo abordarías la implementación de un SGSI ISO 27001 desde cero?",
    "siem": "¿Qué reglas de correlación priorizarías en un SIEM y por qué?",
    "firewall": "Explica cómo segmentarías una red con firewalls y qué reglas base aplicarías.",
    "respuesta a incidentes": "Describe tu proceso de respuesta ante un incidente de ransomware.",
    "pentesting": "¿Qué fases sigue una prueba de intrusión y qué herramientas usas en cada una?",
    "hardening": "¿Qué medidas de hardening aplicarías a un servidor Windows/Linux?",
    "vulnerabilidades": "¿Cómo priorizas la remediación de vulnerabilidades (CVSS, exposición)?",
    "active directory": "¿Cómo asegurarías un dominio de Active Directory?",
    "redes": "Explica el modelo TCP/IP y cómo diagnosticarías un problema de conectividad.",
    "soporte tecnico": "Cuéntame cómo resolviste una incidencia complicada de un usuario.",
    "itil": "¿Cómo aplicas la gestión de incidencias y problemas según ITIL?",
    "cloud": "¿Qué consideraciones de seguridad tienes en un entorno cloud (AWS/Azure)?",
    "python": "¿Para qué has automatizado tareas con Python?",
    "sql": "Escribe una consulta para obtener los registros duplicados de una tabla.",
    "react": "¿Cómo gestionas el estado en una app React grande?",
}


def interview_prep(role, matched, missing):
    """Genera preguntas probables + checklist de preparación."""
    role = role or "el puesto"
    questions = [
        f"Cuéntame sobre ti y por qué te interesa {role}.",
        "¿Por qué quieres trabajar en esta empresa?",
        "Describe un logro del que estés orgulloso/a en tu carrera.",
        "Cuéntame de un reto técnico difícil y cómo lo resolviste.",
        "¿Cómo te mantienes actualizado/a en tu área?",
    ]
    for sk in matched[:5]:
        if sk in SKILL_QUESTIONS:
            questions.append(SKILL_QUESTIONS[sk])
    for sk in missing[:3]:
        questions.append(f"Puede que te pregunten por «{sk}» (aparece en la vacante y no está en tu CV): prepara una respuesta honesta.")
    questions.append("¿Tienes alguna pregunta para nosotros? (prepara 2-3).")

    checklist = [
        "Investiga la empresa: qué hacen, noticias recientes y su cultura.",
        "Repasa la descripción de la vacante y relaciona cada requisito con un ejemplo tuyo.",
        "Prepara 2-3 historias con el método STAR (Situación, Tarea, Acción, Resultado).",
        "Ten a mano tu CV y ejemplos concretos de tus logros con números.",
        "Prepara preguntas inteligentes para el entrevistador.",
        "Comprueba la parte técnica: cámara, micrófono y conexión si es remota.",
        "Prepara tu respuesta a «¿cuáles son tus expectativas salariales?».",
    ]
    if missing:
        checklist.append("Repasa por encima: " + ", ".join(missing[:5]) +
                         " (aparecen en la vacante y son tus brechas).")
    return {"questions": questions, "checklist": checklist}


def requirements_breakdown(job_text, cv_text, sd):
    """Extrae 'requisitos vs lo que tienes' a partir de las habilidades y del texto."""
    job_sk = detect_skills(job_text, sd)
    cv_sk = detect_skills(cv_text, sd)
    rows = []
    for sk in sorted(job_sk):
        rows.append({"req": sk, "have": sk in cv_sk})
    # Años de experiencia solicitados
    m = re.search(r"(\d+)\+?\s*(?:años|years|anos)", job_text.lower())
    if m:
        rows.append({"req": f"{m.group(1)}+ años de experiencia", "have": None})
    # Idiomas
    for lang in ["inglés", "ingles", "english", "portugués", "portugues"]:
        if lang in job_text.lower():
            rows.append({"req": f"Idioma: {lang}", "have": lang in cv_text.lower()})
            break
    have = sum(1 for r in rows if r["have"] is True)
    total = sum(1 for r in rows if r["have"] is not None)
    return {"rows": rows, "have": have, "total": total,
            "pct": round(have / total * 100) if total else 0}


def build_digest_text(user, conn, freq="daily"):
    """Resumen de actividad reciente para correo/Telegram."""
    days = 1 if freq == "daily" else 7
    since = (datetime.now() - timedelta(days=days)).isoformat()
    since_d = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
    found = conn.execute("SELECT COUNT(*) n FROM jobs WHERE user=? AND found_at>=?",
                         (user, since)).fetchone()["n"]
    matches = conn.execute("SELECT COUNT(*) n FROM alerts WHERE user=? AND created_at>=?",
                           (user, since)).fetchone()["n"]
    contacted = conn.execute("SELECT COUNT(*) n FROM outreach WHERE user=? AND fecha>=?",
                             (user, since_d)).fetchone()["n"]
    top = conn.execute("""SELECT title,company,score FROM jobs WHERE user=? AND found_at>=?
                       AND score IS NOT NULL ORDER BY score DESC LIMIT 3""",
                       (user, since)).fetchall()
    label = "hoy" if freq == "daily" else "esta semana"
    lines = [f"📊 Resumen de {label}:", "",
             f"• {found} vacante(s) nuevas", f"• {matches} coincidieron con tu perfil",
             f"• {contacted} empresa(s) contactadas"]
    if top:
        lines.append("")
        lines.append("Mejores encajes:")
        for t in top:
            lines.append(f"  - {t['title']} · {t['company']} ({t['score']:.0f}/100)")
    return "\n".join(lines)


# ============================================================================
#  CV ATS (.docx)
# ============================================================================

def build_ats_cv(profile, cv_text, job, matched, extra_keywords, out_path):
    import docx
    from docx.shared import Pt, RGBColor
    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    h = doc.add_paragraph()
    r = h.add_run(profile.get("profile_nombre") or "Nombre Apellido")
    r.bold = True
    r.font.size = Pt(20)
    sub = doc.add_paragraph()
    sr = sub.add_run(profile.get("profile_titulo") or "")
    sr.font.size = Pt(12)
    sr.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_heading("Perfil profesional", level=1)
    resumen_kw = ", ".join(matched[:8]) if matched else (profile.get("profile_desc_corto") or "")
    doc.add_paragraph(
        f"{profile.get('profile_titulo','')} con {profile.get('profile_desc_corto','')}. "
        f"Competencias alineadas con la vacante de {job.get('title','el puesto')}: "
        f"{resumen_kw}.")

    doc.add_heading("Competencias clave", level=1)
    all_kw = list(dict.fromkeys(matched + (extra_keywords or [])))
    if all_kw:
        for kw in all_kw:
            doc.add_paragraph(kw.capitalize(), style="List Bullet")
    else:
        doc.add_paragraph("(Analiza una vacante para poblar esta sección.)")

    doc.add_heading("Experiencia y formación (de tu CV)", level=1)
    if cv_text.strip():
        for line in cv_text.splitlines():
            if line.strip():
                doc.add_paragraph(line.rstrip())
    else:
        doc.add_paragraph("(No hay CV cargado.)")

    note = doc.add_paragraph()
    nr = note.add_run("Documento generado como base optimizada para ATS. Revisa el "
                      "formato final y asegúrate de que cada competencia refleje "
                      "experiencia real.")
    nr.italic = True
    nr.font.size = Pt(8)
    nr.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.save(out_path)
    return out_path


def _docx_all_text(doc):
    """Texto completo del docx, incluyendo tablas."""
    parts = [p.text for p in doc.paragraphs]
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def _match_font(paragraph, run):
    """Copia tamaño/fuente del primer run existente del párrafo al nuevo run."""
    for r in paragraph.runs:
        if r is run:
            continue
        try:
            if r.font.size:
                run.font.size = r.font.size
            if r.font.name:
                run.font.name = r.font.name
        except Exception:
            pass
        break


SKILL_HINTS = ("habilidad", "competenc", "skill", "conocimiento", "aptitud",
               "tecnolog", "herramient", "aptitudes", "expertise")


def build_ats_cv_preserve(original_path, keywords, out_path):
    """
    EDITA el .docx original conservando su presentación (fuentes, colores,
    maquetación). Inserta solo las keywords que faltan dentro de la sección de
    habilidades existente (o, si no la hay, en una línea compacta al final).
    Así se mantiene el mismo número de páginas en la práctica.
    Devuelve dict con lo añadido y dónde.
    """
    import docx
    doc = docx.Document(original_path)
    full = _docx_all_text(doc).lower()
    to_add = [k.strip() for k in keywords if k.strip() and k.strip().lower() not in full]
    if not to_add:
        doc.save(out_path)
        return {"added": [], "where": "nada", "preserved": True,
                "note": "Todas las keywords seleccionadas ya estaban en tu CV."}
    added_str = ", ".join(x.capitalize() for x in to_add)

    heading = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip().lower()
        if t and len(t) < 45 and any(h in t for h in SKILL_HINTS):
            heading = i
            break

    if heading is not None:
        insert = None
        for j in range(heading + 1, min(heading + 4, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                insert = doc.paragraphs[j]
                break
        if insert is None:
            insert = doc.paragraphs[heading]
        tail = insert.text.strip()
        sep = "" if (not tail or tail[-1] in ",;:·•-") else ", "
        run = insert.add_run(sep + added_str)
        _match_font(insert, run)
        where = "tu sección de habilidades"
    else:
        p = doc.add_paragraph()
        r = p.add_run("Competencias adicionales: " + added_str)
        try:
            r.font.size = doc.styles["Normal"].font.size
        except Exception:
            pass
        where = "una línea al final del documento"

    doc.save(out_path)
    return {"added": to_add, "where": where, "preserved": True}


def profile_scoring_text(user, cv_text=None, conn=None):
    """Texto para puntuar: CV + perfil + área de interés del usuario."""
    cv = cv_text if cv_text is not None else get_config(user, "cv_text", conn)
    pr = user_profile(user, conn)
    return " ".join([cv or "", pr.get("profile_titulo", ""),
                     pr.get("profile_desc_corto", ""), pr.get("profile_desc_largo", "")])


def all_user_cvs(user, conn=None):
    """Lista de CVs del usuario: el principal + las variantes guardadas."""
    conn = conn or db()
    out = []
    primary = get_config(user, "cv_text", conn)
    if primary and primary.strip():
        ext = get_config(user, "cv_ext", conn)
        out.append({"id": 0, "label": "Principal", "text": primary, "ext": ext,
                    "orig": os.path.join(ORIG_DIR, f"{user}{ext}")})
    for r in conn.execute("SELECT id,label,text,ext FROM cv_variants WHERE user=?",
                          (user,)).fetchall():
        out.append({"id": r["id"], "label": r["label"], "text": r["text"],
                    "ext": r["ext"], "orig": os.path.join(ORIG_DIR, f"{user}__{r['id']}{r['ext']}")})
    return out


def best_score(user, job_text, sd, conn=None):
    """Puntúa la vacante contra TODOS los CVs y devuelve el mejor (score, label)."""
    cvs = all_user_cvs(user, conn)
    if not cvs:
        return None, None
    pr = user_profile(user, conn)
    prof_suffix = " " + " ".join([pr.get("profile_titulo", ""),
                                  pr.get("profile_desc_corto", ""),
                                  pr.get("profile_desc_largo", "")])
    best, best_label = None, None
    for cv in cvs:
        sc = score_cv(cv["text"] + prof_suffix, job_text, sd)
        if sc and (best is None or sc["score"] > best["score"]):
            best, best_label = sc, cv["label"]
    return best, best_label


def cv_by_label(user, label, conn=None):
    for cv in all_user_cvs(user, conn):
        if cv["label"] == label:
            return cv
    cvs = all_user_cvs(user, conn)
    return cvs[0] if cvs else None


# ============================================================================
#  AUTENTICACIÓN (users.txt con hash)
# ============================================================================

_users_lock = threading.Lock()
USER_RE = re.compile(r"^[A-Za-z0-9_.\-]{3,32}$")


def _hash_pw(password, salt=None):
    if salt is None:
        salt = os.urandom(16)
    dk = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt, 120000)
    return binascii.hexlify(dk).decode(), binascii.hexlify(salt).decode()


def load_users():
    users = {}
    if not os.path.exists(USERS_TXT):
        return users
    with open(USERS_TXT, "r", encoding="utf-8") as f:
        for line in f:
            parts = line.rstrip("\n").split("\t")
            if len(parts) >= 3:
                users[parts[0]] = {"hash": parts[1], "salt": parts[2],
                                   "created": parts[3] if len(parts) > 3 else ""}
    return users


def register_user(username, password):
    username = (username or "").strip()
    if not USER_RE.match(username):
        return False, "Usuario inválido (3-32 caracteres: letras, números, . _ -)."
    if len(password or "") < 4:
        return False, "La contraseña debe tener al menos 4 caracteres."
    with _users_lock:
        if username in load_users():
            return False, "Ese usuario ya existe."
        h, salt = _hash_pw(password)
        with open(USERS_TXT, "a", encoding="utf-8") as f:
            f.write(f"{username}\t{h}\t{salt}\t{datetime.now().isoformat()}\n")
    seed_user_config(username)
    return True, "ok"


def verify_user(username, password):
    u = load_users().get((username or "").strip())
    if not u:
        return False
    h, _ = _hash_pw(password, binascii.unhexlify(u["salt"]))
    return h == u["hash"]


# ---- Cuentas, roles y verificación (multiusuario) --------------------------
OWNER_USERNAME = "0xJohnnyHash"      # dueño/CEO: verificado y con rol owner automáticamente
ROLES = ("user", "company", "moderator", "admin", "owner")
# Roles asignables por un admin (el owner además puede asignar admin)
ASSIGNABLE_ROLES = ("user", "company", "moderator", "admin")


def ensure_account(username):
    """Crea la fila de cuenta si no existe y aplica el rol de dueño al OWNER."""
    c = db()
    row = c.execute("SELECT * FROM accounts WHERE username=?", (username,)).fetchone()
    if not row:
        is_owner = username.lower() == OWNER_USERNAME.lower()
        c.execute("INSERT INTO accounts(username,role,verified,display_name,created_at) VALUES(?,?,?,?,?)",
                  (username, "owner" if is_owner else "user", 1 if is_owner else 0,
                   username, datetime.now().isoformat()))
        c.commit()
        row = c.execute("SELECT * FROM accounts WHERE username=?", (username,)).fetchone()
    elif username.lower() == OWNER_USERNAME.lower() and (row["role"] != "owner" or row["verified"] != 1):
        c.execute("UPDATE accounts SET role='owner', verified=1 WHERE username=?", (username,))
        c.commit()
        row = c.execute("SELECT * FROM accounts WHERE username=?", (username,)).fetchone()
    return dict(row)


def get_account(username):
    return ensure_account(username)


def is_admin(username):
    return get_account(username).get("role") in ("admin", "owner")


def is_moderator(username):
    return get_account(username).get("role") in ("moderator", "admin", "owner")


def is_verified(username):
    a = get_account(username)
    return a.get("verified") == 1 or a.get("role") in ("moderator", "admin", "owner")


def account_public(username):
    """Datos públicos de una cuenta para mostrar en chat/bolsa/perfil."""
    a = get_account(username)
    return {"username": username, "role": a.get("role"), "verified": 1 if is_verified(username) else 0,
            "display_name": a.get("display_name") or username, "headline": a.get("headline") or "",
            "bio": a.get("bio") or "", "location": a.get("location") or "", "links": a.get("links") or "",
            "has_avatar": bool(a.get("avatar_ext")), "avatar_ext": a.get("avatar_ext") or "",
            "check_color": a.get("check_color") or ""}


def update_password(username, old, new):
    if not verify_user(username, old):
        return False, "La contraseña actual no es correcta."
    if len(new or "") < 4:
        return False, "La nueva contraseña debe tener al menos 4 caracteres."
    with _users_lock:
        users = load_users()
        if username not in users:
            return False, "El usuario no existe."
        h, salt = _hash_pw(new)
        lines = []
        for uname, info in users.items():
            if uname == username:
                lines.append(f"{uname}\t{h}\t{salt}\t{info.get('created', '')}")
            else:
                lines.append(f"{uname}\t{info['hash']}\t{info['salt']}\t{info.get('created', '')}")
        with open(USERS_TXT, "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")
    return True, "ok"


# ============================================================================
#  BASE DE DATOS (por usuario)
# ============================================================================

_local = threading.local()


def db():
    if not hasattr(_local, "conn"):
        _local.conn = sqlite3.connect(DB_PATH)
        _local.conn.row_factory = sqlite3.Row
        _local.conn.execute("PRAGMA journal_mode=WAL")
    return _local.conn


def init_db():
    c = db()
    c.executescript("""
    CREATE TABLE IF NOT EXISTS config (
        user TEXT, key TEXT, value TEXT, PRIMARY KEY(user,key));
    CREATE TABLE IF NOT EXISTS searches (
        id INTEGER PRIMARY KEY AUTOINCREMENT, user TEXT,
        keyword TEXT, country TEXT, city TEXT, active INTEGER DEFAULT 1,
        created_at TEXT, last_run TEXT);
    CREATE TABLE IF NOT EXISTS jobs (
        user TEXT, id TEXT, title TEXT, company TEXT, location TEXT,
        url TEXT, date TEXT, description TEXT, easy_apply INTEGER,
        score REAL, verdict TEXT, matched TEXT, missing TEXT,
        source_query TEXT, found_at TEXT, seen INTEGER DEFAULT 0,
        PRIMARY KEY(user,id));
    CREATE TABLE IF NOT EXISTS alerts (
        id INTEGER PRIMARY KEY AUTOINCREMENT, user TEXT, job_id TEXT,
        score REAL, created_at TEXT, read INTEGER DEFAULT 0);
    CREATE TABLE IF NOT EXISTS outreach (
        id INTEGER PRIMARY KEY AUTOINCREMENT, user TEXT, empresa TEXT, rol TEXT,
        contacto TEXT, url TEXT, job_id TEXT, notas TEXT, fecha TEXT);
    CREATE TABLE IF NOT EXISTS gen_cvs (
        user TEXT, job_id TEXT, title TEXT, company TEXT, path TEXT,
        before REAL, after REAL, created_at TEXT, PRIMARY KEY(user,job_id));
    CREATE TABLE IF NOT EXISTS cv_variants (
        id INTEGER PRIMARY KEY AUTOINCREMENT, user TEXT, label TEXT,
        text TEXT, ext TEXT, filename TEXT, created TEXT);
    CREATE TABLE IF NOT EXISTS accounts (
        username TEXT PRIMARY KEY, role TEXT DEFAULT 'user', verified INTEGER DEFAULT 0,
        display_name TEXT, headline TEXT, bio TEXT, location TEXT, links TEXT,
        avatar_ext TEXT, created_at TEXT);
    CREATE TABLE IF NOT EXISTS rooms (
        id INTEGER PRIMARY KEY AUTOINCREMENT, name TEXT UNIQUE, created_by TEXT, created_at TEXT);
    CREATE TABLE IF NOT EXISTS chat (
        id INTEGER PRIMARY KEY AUTOINCREMENT, username TEXT, body TEXT,
        created_at TEXT, deleted INTEGER DEFAULT 0);
    CREATE TABLE IF NOT EXISTS presence (
        username TEXT PRIMARY KEY, last_seen TEXT);
    CREATE TABLE IF NOT EXISTS pms (
        id INTEGER PRIMARY KEY AUTOINCREMENT, sender TEXT, recipient TEXT, body TEXT,
        created_at TEXT, read INTEGER DEFAULT 0);
    CREATE TABLE IF NOT EXISTS friends (
        requester TEXT, addressee TEXT, status TEXT DEFAULT 'pending', created_at TEXT,
        PRIMARY KEY(requester, addressee));
    CREATE TABLE IF NOT EXISTS listings (
        id INTEGER PRIMARY KEY AUTOINCREMENT, company TEXT, title TEXT, description TEXT,
        location TEXT, remote INTEGER DEFAULT 0, salary TEXT, skills TEXT,
        created_at TEXT, active INTEGER DEFAULT 1);
    CREATE TABLE IF NOT EXISTS applications (
        id INTEGER PRIMARY KEY AUTOINCREMENT, listing_id INTEGER, applicant TEXT,
        message TEXT, created_at TEXT, UNIQUE(listing_id, applicant));
    """)
    # Migraciones suaves de columnas nuevas
    for tbl, col, decl in [("jobs", "source", "TEXT"), ("jobs", "app_status", "TEXT"),
                           ("jobs", "best_cv", "TEXT"),
                           ("jobs", "favorite", "INTEGER DEFAULT 0"),
                           ("jobs", "interview_at", "TEXT"),
                           ("jobs", "tags", "TEXT"),
                           ("jobs", "archived", "INTEGER DEFAULT 0"),
                           ("jobs", "notes", "TEXT"),
                           ("jobs", "applied_at", "TEXT"),
                           ("outreach", "responded", "INTEGER DEFAULT 0"),
                           ("outreach", "template", "TEXT"),
                           ("chat", "room_id", "INTEGER DEFAULT 1"),
                           ("accounts", "check_color", "TEXT")]:
        try:
            c.execute(f"ALTER TABLE {tbl} ADD COLUMN {col} {decl}")
        except Exception:
            pass
    # Salas de chat por defecto: #general + países (id 1 = general para chats antiguos)
    if not c.execute("SELECT 1 FROM rooms LIMIT 1").fetchone():
        default_rooms = ["general", "venezuela", "colombia", "peru", "chile", "argentina",
                         "mexico", "ecuador", "espana", "estados-unidos", "remoto"]
        for name in default_rooms:
            c.execute("INSERT OR IGNORE INTO rooms(name,created_by,created_at) VALUES(?,?,?)",
                      (name, OWNER_USERNAME, datetime.now().isoformat()))
    c.commit()


def seed_user_config(user):
    c = db()
    for k, v in DEFAULTS.items():
        c.execute("INSERT OR IGNORE INTO config(user,key,value) VALUES(?,?,?)",
                  (user, k, v))
    c.commit()


def get_config(user, key, conn=None):
    conn = conn or db()
    row = conn.execute("SELECT value FROM config WHERE user=? AND key=?",
                       (user, key)).fetchone()
    return row["value"] if row is not None else DEFAULTS.get(key, "")


def set_config(user, key, value, conn=None):
    conn = conn or db()
    conn.execute("INSERT INTO config(user,key,value) VALUES(?,?,?) "
                 "ON CONFLICT(user,key) DO UPDATE SET value=excluded.value",
                 (user, key, str(value)))
    conn.commit()


def user_profile(user, conn=None):
    return {k: get_config(user, k, conn) for k in
            ("profile_nombre", "profile_titulo", "profile_desc_corto",
             "profile_desc_largo")}


def user_skills(user, conn=None):
    return parse_skill_dict(get_config(user, "skill_dict", conn))


def top_cv_skills(user, conn=None):
    cv = get_config(user, "cv_text", conn)
    return sorted(detect_skills(cv, user_skills(user, conn))) if cv else []


# ============================================================================
#  MOTOR DE ESCANEO EN SEGUNDO PLANO (multiusuario)
# ============================================================================

def _dedup_key(title, company):
    """Clave normalizada para detectar la misma vacante entre portales."""
    norm = lambda s: re.sub(r"[^a-z0-9]+", "", (s or "").lower())
    return norm(title) + "|" + norm(company)


class Scanner(threading.Thread):
    def __init__(self):
        super().__init__(daemon=True)
        self.wake = threading.Event()
        self.status = {}       # por usuario
        self.scan_count = {}   # nº de escaneos completados por usuario (para animar el logo)
        self.running = False

    def run(self):
        time.sleep(3)
        while True:
            try:
                self.scan_active_users()
            except Exception as e:
                print("Scanner error:", e)
            # Intervalo = el más corto entre los usuarios activos (mín 5 min)
            self.wake.wait(timeout=self._min_interval() * 60)
            self.wake.clear()

    def trigger(self):
        self.wake.set()

    def _min_interval(self):
        conn = _tc()
        vals = [int(r["value"] or 30) for r in conn.execute(
            "SELECT value FROM config WHERE key='scan_interval_min'").fetchall()]
        conn.close()
        return max(5, min(vals) if vals else 30)

    def scan_active_users(self):
        conn = _tc()
        users = [r["user"] for r in conn.execute(
            "SELECT DISTINCT user FROM searches WHERE active=1").fetchall()]
        # Solo escanea usuarios con escaneo activado
        result = {}
        for u in users:
            enabled = get_config(u, "scan_enabled", conn) == "1"
            if enabled:
                result[u] = None
        conn.close()
        for u in result:
            try:
                self.scan_user(u)
            except Exception as e:
                self.status[u] = f"Error: {e}"

    def scan_user(self, user, force=False):
        conn = _tc()
        if not force and get_config(user, "scan_enabled", conn) != "1":
            conn.close()
            return {"new": 0, "matches": 0}
        rows = conn.execute("SELECT * FROM searches WHERE user=? AND active=1",
                            (user,)).fetchall()
        if not rows:
            self.status[user] = "Sin búsquedas activas"
            conn.close()
            return {"new": 0, "matches": 0}
        sd = user_skills(user, conn)
        has_cv = bool(all_user_cvs(user, conn))
        ecfg = email_config(user, conn)
        threshold = float(get_config(user, "match_threshold", conn) or 60)
        maxr = int(get_config(user, "max_per_search", conn) or 25)
        rss = [x for x in (get_config(user, "rss_feeds", conn) or "").splitlines() if x.strip()]
        existing = {r["id"] for r in conn.execute(
            "SELECT id FROM jobs WHERE user=?", (user,)).fetchall()}
        dedup = {_dedup_key(r["title"], r["company"]) for r in conn.execute(
            "SELECT title,company FROM jobs WHERE user=?", (user,)).fetchall()}
        tcfg = {k: get_config(user, k, conn) for k in
                ("telegram_enabled", "telegram_token", "telegram_chat_id")}
        srcs = [x.strip() for x in (get_config(user, "sources", conn) or "").split(",") if x.strip()]
        adz = {"app_id": get_config(user, "adzuna_app_id", conn),
               "app_key": get_config(user, "adzuna_app_key", conn),
               "country": get_config(user, "adzuna_country", conn) or "es"}
        self.running = True
        total_new, total_match = 0, 0

        def process(job, source, source_query, need_detail):
            nonlocal total_new, total_match
            if job["id"] in existing:
                return
            dk = _dedup_key(job["title"], job["company"])
            if dk in dedup:
                return
            if need_detail:
                desc, easy = LI.details(job["id"])
                job["description"] = desc
                job["easy_apply"] = easy
                time.sleep(random.uniform(0.6, 1.2))
            desc = job.get("description", "")
            job_text = f"{job['title']} {desc}"
            sc, bestcv = best_score(user, job_text, sd, conn) if has_cv else (None, None)
            conn.execute("""INSERT OR IGNORE INTO jobs(user,id,title,company,location,
                url,date,description,easy_apply,score,verdict,matched,missing,
                source_query,found_at,seen,source,app_status,best_cv)
                VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,0,?,?,?)""",
                (user, job["id"], job["title"], job["company"], job.get("location", ""),
                 job.get("url", ""), job.get("date", ""), desc, job.get("easy_apply"),
                 sc["score"] if sc else None, sc["verdict"] if sc else None,
                 ", ".join(sc["matched"]) if sc else "",
                 ", ".join(sc["missing"]) if sc else "",
                 source_query, datetime.now().isoformat(), source, "new", bestcv))
            conn.commit()
            existing.add(job["id"])
            dedup.add(dk)
            total_new += 1
            if sc and sc["score"] >= threshold:
                conn.execute("INSERT INTO alerts(user,job_id,score,created_at,read) "
                             "VALUES(?,?,?,?,0)",
                             (user, job["id"], sc["score"], datetime.now().isoformat()))
                conn.commit()
                total_match += 1
                portal = SOURCE_NAMES.get(source, source)
                notify_desktop("¡Vacante que encaja!",
                               f"{job['title']} — {job['company']} ({sc['score']:.0f}/100)")
                if ecfg.get("email_enabled") == "1":
                    try:
                        send_email(ecfg, f"[Vacantes] {job['title']} — {sc['score']:.0f}/100",
                                   f"Nueva vacante que encaja con tu perfil:\n\n"
                                   f"{job['title']}\n{job['company']} · {job.get('location','')}\n"
                                   f"Encaje: {sc['score']:.0f}/100 · Portal: {portal}\n\n"
                                   f"{job.get('url','')}\n\n— Buscador de Vacantes")
                    except Exception:
                        pass
                if tcfg.get("telegram_enabled") == "1":
                    try:
                        send_telegram(tcfg["telegram_token"], tcfg["telegram_chat_id"],
                                      f"🎯 Vacante que encaja ({sc['score']:.0f}/100)\n"
                                      f"{job['title']}\n{job['company']} · {job.get('location','')}\n"
                                      f"Portal: {portal}\n{job.get('url','')}")
                    except Exception:
                        pass

        for s in rows:
            loc = f"{s['city']}, {s['country']}" if s["city"] else s["country"]
            sq = f"{s['keyword']} @ {loc}"
            self.status[user] = f"Escaneando «{s['keyword']}»..."
            try:
                li = LI.search(s["keyword"], loc, max_results=maxr)
            except Exception:
                li = []
            for job in li:
                process(job, "linkedin", sq, True)
            portal_jobs = fetch_sources(
                srcs, s["keyword"], loc, maxr,
                adzuna=adz if "adzuna" in srcs else None,
                rss=rss if "rss" in srcs else None,
                log=lambda m: self.status.__setitem__(user, m))
            for job in portal_jobs:
                process(job, job.get("source", "portal"), sq, False)
            conn.execute("UPDATE searches SET last_run=? WHERE id=?",
                         (datetime.now().isoformat(), s["id"]))
            conn.commit()
        self.running = False
        self.scan_count[user] = self.scan_count.get(user, 0) + 1
        self.status[user] = (f"Último escaneo: {datetime.now():%H:%M} · "
                             f"{total_new} nuevas, {total_match} coincidencias")
        try:
            self._maybe_digest(user, conn, ecfg)
        except Exception:
            pass
        conn.close()
        return {"new": total_new, "matches": total_match}

    def _maybe_digest(self, user, conn, ecfg):
        if get_config(user, "digest_enabled", conn) != "1":
            return
        freq = get_config(user, "digest_freq", conn) or "daily"
        last = get_config(user, "last_digest", conn)
        today = datetime.now().date()
        if last:
            try:
                gap = (today - datetime.fromisoformat(last).date()).days
                if (freq == "daily" and gap < 1) or (freq == "weekly" and gap < 7):
                    return
            except Exception:
                pass
        text = build_digest_text(user, conn, freq)
        chan = get_config(user, "digest_channel", conn) or "email"
        if chan in ("email", "both") and ecfg.get("email_enabled") == "1":
            try:
                send_email(ecfg, "[Vacantes] Tu resumen", text)
            except Exception:
                pass
        if chan in ("telegram", "both") and get_config(user, "telegram_enabled", conn) == "1":
            try:
                send_telegram(get_config(user, "telegram_token", conn),
                              get_config(user, "telegram_chat_id", conn), text)
            except Exception:
                pass
        set_config(user, "last_digest", datetime.now().isoformat(), conn)


def _tc():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def send_email(cfg, subject, body):
    """Envía un correo usando la config SMTP del usuario (dict). Lanza excepción
    si falla (para poder reportar en el correo de prueba)."""
    host = cfg.get("smtp_host", "").strip()
    port = int(cfg.get("smtp_port") or 587)
    user = cfg.get("smtp_user", "").strip()
    pwd = cfg.get("smtp_pass", "")
    to = cfg.get("email_to", "").strip() or user
    if not (host and user and pwd and to):
        raise RuntimeError("Faltan datos SMTP (host, usuario, contraseña o destinatario).")
    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = subject
    msg["From"] = user
    msg["To"] = to
    if port == 465:
        with smtplib.SMTP_SSL(host, port, timeout=20) as s:
            s.login(user, pwd)
            s.sendmail(user, [to], msg.as_string())
    else:
        with smtplib.SMTP(host, port, timeout=20) as s:
            s.starttls()
            s.login(user, pwd)
            s.sendmail(user, [to], msg.as_string())


def email_config(user, conn=None):
    return {k: get_config(user, k, conn) for k in
            ("email_enabled", "email_to", "smtp_host", "smtp_port",
             "smtp_user", "smtp_pass")}


def notify_desktop(title, message):
    """Notificación de escritorio con varios respaldos según el sistema."""
    ok = False
    try:
        from plyer import notification
        notification.notify(title=title, message=message,
                            app_name="Vacantes", timeout=12)
        ok = True
    except Exception:
        ok = False
    if not ok and sys.platform.startswith("win"):
        # Respaldo Windows: notificación por PowerShell (globo del área de notificación)
        try:
            import subprocess
            ps = (
                "[void][reflection.assembly]::LoadWithPartialName('System.Windows.Forms');"
                "$n=New-Object System.Windows.Forms.NotifyIcon;"
                "$n.Icon=[System.Drawing.SystemIcons]::Information;$n.Visible=$true;"
                f"$n.ShowBalloonTip(12000,'{title}','{message}',"
                "[System.Windows.Forms.ToolTipIcon]::Info);Start-Sleep -Seconds 8;$n.Dispose()")
            subprocess.Popen(["powershell", "-NoProfile", "-WindowStyle", "Hidden",
                              "-Command", ps])
            ok = True
        except Exception:
            ok = False
    if not ok:
        # Último respaldo: un sonido para que te enteres
        try:
            if sys.platform.startswith("win"):
                import winsound
                winsound.MessageBeep()
            else:
                print("\a", end="", flush=True)
        except Exception:
            pass


SCANNER = Scanner()


# ============================================================================
#  APP FLASK
# ============================================================================

app = Flask(__name__, template_folder=os.path.join(APP_DIR, "templates"))


def _get_secret():
    if os.path.exists(SECRET_FILE):
        return open(SECRET_FILE, "rb").read()
    s = os.urandom(24)
    open(SECRET_FILE, "wb").write(s)
    return s


app.secret_key = _get_secret()


def login_required(f):
    @wraps(f)
    def w(*a, **k):
        if not session.get("user"):
            return jsonify({"error": "No has iniciado sesión."}), 401
        return f(*a, **k)
    return w


def cu():
    return session["user"]


# ---- Auth ----
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/me")
def api_me():
    if session.get("user"):
        return jsonify({"user": session["user"]})
    return jsonify({"user": None})


@app.route("/api/register", methods=["POST"])
def api_register():
    d = request.json or {}
    ok, msg = register_user(d.get("username", ""), d.get("password", ""))
    if not ok:
        return jsonify({"error": msg}), 400
    session["user"] = d["username"].strip()
    ensure_account(session["user"])
    # Perfil opcional capturado en el registro
    if d.get("nombre"):
        set_config(session["user"], "profile_nombre", d["nombre"].strip())
        db().execute("UPDATE accounts SET display_name=? WHERE username=?",
                     (d["nombre"].strip(), session["user"])); db().commit()
    if d.get("titulo"):
        set_config(session["user"], "profile_titulo", d["titulo"].strip())
    return jsonify({"ok": True, "user": session["user"]})


@app.route("/api/login", methods=["POST"])
def api_login():
    d = request.json or {}
    if verify_user(d.get("username", ""), d.get("password", "")):
        session["user"] = d["username"].strip()
        seed_user_config(session["user"])
        ensure_account(session["user"])
        return jsonify({"ok": True, "user": session["user"]})
    return jsonify({"error": "Usuario o contraseña incorrectos."}), 401


@app.route("/api/logout", methods=["POST"])
def api_logout():
    session.pop("user", None)
    return jsonify({"ok": True})


# ---- Estado ----
@app.route("/api/state")
@login_required
def api_state():
    u = cu()
    c = db()
    cfg = {r["key"]: r["value"] for r in
           c.execute("SELECT key,value FROM config WHERE user=?", (u,))}
    for k, v in DEFAULTS.items():
        cfg.setdefault(k, v)
    cfg.pop("cv_text", None)
    cfg.pop("scratchpad", None)   # se cargan bajo demanda vía /api/board
    cfg.pop("board_png", None)
    for secret in ("smtp_pass", "telegram_token", "adzuna_app_key", "ai_api_key"):
        cfg[secret + "_set"] = "1" if cfg.get(secret) else "0"
        cfg.pop(secret, None)
    q = lambda sql: c.execute(sql, (u,)).fetchone()["n"]
    counts = {
        "jobs": q("SELECT COUNT(*) n FROM jobs WHERE user=?"),
        "new": q("SELECT COUNT(*) n FROM jobs WHERE user=? AND seen=0"),
        "alerts": q("SELECT COUNT(*) n FROM alerts WHERE user=? AND read=0"),
        "searches": q("SELECT COUNT(*) n FROM searches WHERE user=? AND active=1"),
        "outreach": q("SELECT COUNT(*) n FROM outreach WHERE user=?"),
        "queued": q("SELECT COUNT(*) n FROM jobs WHERE user=? AND app_status='queued'"),
        "pipeline": q("SELECT COUNT(*) n FROM jobs WHERE user=? AND app_status IN "
                      "('queued','applied','interview','offer','rejected')"),
        "favorites": q("SELECT COUNT(*) n FROM jobs WHERE user=? AND favorite=1"),
        "archived": q("SELECT COUNT(*) n FROM jobs WHERE user=? AND archived=1"),
        "applied_week": db().execute(
            "SELECT COUNT(*) n FROM jobs WHERE user=? AND applied_at>=?",
            (u, (datetime.now() - timedelta(days=7)).isoformat())).fetchone()["n"],
        "upcoming": db().execute(
            "SELECT COUNT(*) n FROM jobs WHERE user=? AND interview_at IS NOT NULL "
            "AND interview_at!='' AND interview_at>=?",
            (u, datetime.now().strftime("%Y-%m-%dT00:00"))).fetchone()["n"],
        "followups": db().execute(
            "SELECT COUNT(*) n FROM outreach WHERE user=? AND responded=0 AND fecha<=?",
            (u, (datetime.now() - timedelta(days=int(get_config(u, 'followup_days') or 5))).strftime("%Y-%m-%d %H:%M"))
        ).fetchone()["n"],
    }
    acc = get_account(u)
    counts["pm_unread"] = db().execute(
        "SELECT COUNT(*) n FROM pms WHERE recipient=? AND read=0", (u,)).fetchone()["n"]
    counts["friend_reqs"] = db().execute(
        "SELECT COUNT(*) n FROM friends WHERE addressee=? AND status='pending'", (u,)).fetchone()["n"]
    me = {"username": u, "role": acc.get("role"), "verified": 1 if is_verified(u) else 0,
          "is_admin": is_admin(u), "is_moderator": is_moderator(u), "has_avatar": bool(acc.get("avatar_ext")),
          "check_color": acc.get("check_color") or "", "can_recolor": is_admin(u) or is_moderator(u),
          "display_name": acc.get("display_name") or u, "headline": acc.get("headline") or "",
          "profile_done": bool(acc.get("headline") or acc.get("bio"))}
    return jsonify({"user": u, "me": me, "config": cfg, "counts": counts,
                    "cv_loaded": bool(get_config(u, "cv_text")),
                    "cv_filename": get_config(u, "cv_filename"),
                    "scanner": {"status": SCANNER.status.get(u, "Inactivo"),
                                "running": SCANNER.running,
                                "scan_count": SCANNER.scan_count.get(u, 0),
                                "enabled": get_config(u, "scan_enabled") == "1"}})


# ---- Búsquedas ----
@app.route("/api/searches", methods=["GET", "POST"])
@login_required
def api_searches():
    u = cu()
    c = db()
    if request.method == "POST":
        d = request.json
        c.execute("INSERT INTO searches(user,keyword,country,city,active,created_at) "
                  "VALUES(?,?,?,?,1,?)",
                  (u, d["keyword"].strip(), d.get("country", "Venezuela").strip(),
                   d.get("city", "").strip(), datetime.now().isoformat()))
        c.commit()
        return jsonify({"ok": True})
    rows = c.execute("SELECT * FROM searches WHERE user=? ORDER BY id DESC", (u,)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/searches/<int:sid>", methods=["DELETE"])
@login_required
def api_search_delete(sid):
    db().execute("DELETE FROM searches WHERE id=? AND user=?", (sid, cu()))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/searches/<int:sid>/toggle", methods=["POST"])
@login_required
def api_search_toggle(sid):
    db().execute("UPDATE searches SET active=1-active WHERE id=? AND user=?", (sid, cu()))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/scan/now", methods=["POST"])
@login_required
def api_scan_now():
    return jsonify(SCANNER.scan_user(cu(), force=True))


# ---- Vacantes ----
# Detección de salario en el texto de la vacante -----------------------------
_SAL_RE = re.compile(
    r"(?:€|\$|£|USD|EUR|GBP)\s?\d[\d.,]*\s?(?:k|mil)?"
    r"(?:\s?[-–a]\s?(?:€|\$|£)?\s?\d[\d.,]*\s?(?:k|mil)?)?"
    r"|\d[\d.,]*\s?(?:€|\$|£|USD|EUR|GBP|k)\b"
    r"|\d[\d.,]*\s?(?:€|\$|£|k)?\s?(?:/\s?(?:año|ano|year|yr|mes|month|hora|hour|h))"
    r"|\d{2,3}[.,]?\d{3}\s?(?:€|\$|£|USD|EUR|GBP)",
    re.IGNORECASE)


def detect_salary(text):
    """Devuelve el primer rango/valor salarial detectado, o '' si no hay."""
    if not text:
        return ""
    m = _SAL_RE.search(text)
    if not m:
        return ""
    s = re.sub(r"\s+", " ", m.group(0)).strip(" .,")
    return s[:40]


@app.route("/api/jobs")
@login_required
def api_jobs():
    u = cu()
    q = "SELECT * FROM jobs WHERE user=?"
    p = [u]
    if request.args.get("new") == "1":
        q += " AND seen=0"
    if request.args.get("easy") == "1":
        q += " AND easy_apply=1"
    if request.args.get("favorite") == "1":
        q += " AND favorite=1"
    ms = request.args.get("min_score", type=float)
    if ms:
        q += " AND score IS NOT NULL AND score>=?"
        p.append(ms)
    src = request.args.get("source")
    if src:
        q += " AND source=?"
        p.append(src)
    # Por defecto se ocultan las archivadas; con archived=1 se muestran SOLO esas
    if request.args.get("archived") == "1":
        q += " AND archived=1"
    else:
        q += " AND (archived IS NULL OR archived=0)"
    # Orden por defecto: NUEVAS (sin ver) primero, luego lo más reciente
    q += " ORDER BY seen ASC, found_at DESC, (score IS NULL), score DESC LIMIT 400"
    jobs = [dict(r) for r in db().execute(q, p).fetchall()]
    sizes = company_postings(u)
    # Filtros calculados
    want_salary = request.args.get("has_salary") == "1"
    days = request.args.get("days", type=int)
    tagf = (request.args.get("tag") or "").strip().lower()
    cutoff = None
    if days:
        cutoff = (datetime.now() - timedelta(days=days)).isoformat()
    for j in jobs:
        n = sizes.get((j["company"] or "").lower(), 1)
        j["company_postings"] = n
        j["size_hint"] = "xs" if n <= 2 else ("sm" if n <= 6 else "lg")
        j["favorite"] = 1 if j.get("favorite") == 1 else 0
        j["archived"] = 1 if j.get("archived") == 1 else 0
        j["is_new"] = 1 if j.get("seen") == 0 else 0   # 🆕 nueva chamba: aún no la has abierto
        j["has_notes"] = 1 if (j.get("notes") or "").strip() else 0
        j["salary"] = detect_salary(j.get("description") or "")
        j["tag_list"] = [t for t in (j.get("tags") or "").split(",") if t.strip()]
        # Ranking inteligente: score + bonus por easy apply y empresa pequeña
        base = j["score"] if j["score"] is not None else 0
        rank = base + (12 if j.get("easy_apply") == 1 else 0) + (10 if j["size_hint"] == "xs" else 0)
        j["rank"] = round(rank, 1)
        j["hot"] = rank >= 85
    if request.args.get("size") == "xs":
        jobs = [j for j in jobs if j["size_hint"] == "xs"]
    if want_salary:
        jobs = [j for j in jobs if j["salary"]]
    if cutoff:
        jobs = [j for j in jobs if (j.get("found_at") or "") >= cutoff]
    if tagf:
        jobs = [j for j in jobs if any(tagf == t.strip().lower() for t in j["tag_list"])]
    if request.args.get("sort") == "smart":
        jobs.sort(key=lambda x: x["rank"], reverse=True)
    return jsonify(jobs)


def company_postings(user, conn=None):
    conn = conn or db()
    rows = conn.execute("SELECT lower(company) c, COUNT(*) n FROM jobs "
                        "WHERE user=? GROUP BY lower(company)", (user,)).fetchall()
    return {r["c"]: r["n"] for r in rows}


@app.route("/api/jobs/<jid>")
@login_required
def api_job_detail(jid):
    u = cu()
    c = db()
    row = c.execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    job = dict(row)
    # Cargar descripción si falta — SOLO para vacantes de LinkedIn
    src = job.get("source") or "linkedin"
    if src == "linkedin" and (not job.get("description") or job["easy_apply"] is None):
        desc, easy = LI.details(jid)
        job["description"] = desc
        job["easy_apply"] = easy
    # PUNTUACIÓN AUTOMÁTICA con el MEJOR CV (perfil + área) al abrir
    sd = user_skills(u)
    sc, bestcv = best_score(u, f"{job['title']} {job['description']}", sd)
    score_by_ai = False
    # Si hay IA configurada, re-evalúa el encaje con IA (offline si no hay key)
    if get_config(u, "ai_enabled") == "1" and get_config(u, "cv_text"):
        try:
            aicfg = {k: get_config(u, k) for k in ("ai_provider", "ai_model", "ai_api_key", "ai_base_url")}
            cvsel = cv_by_label(u, bestcv) if bestcv else None
            cvtext = (cvsel or {}).get("text") or get_config(u, "cv_text")
            ai_sc = ai_score_cv(aicfg, cvtext, job["title"], job["description"])
            if ai_sc:
                sc = ai_sc
                score_by_ai = True
        except Exception:
            pass
    c.execute("""UPDATE jobs SET description=?,easy_apply=?,score=?,verdict=?,
                 matched=?,missing=?,best_cv=?,seen=1 WHERE user=? AND id=?""",
              (job["description"], job["easy_apply"],
               sc["score"] if sc else job["score"], sc["verdict"] if sc else job["verdict"],
               ", ".join(sc["matched"]) if sc else job["matched"],
               ", ".join(sc["missing"]) if sc else job["missing"],
               bestcv if sc else job.get("best_cv"), u, jid))
    c.commit()
    job = dict(c.execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone())
    job["score_by_ai"] = score_by_ai
    job["people"] = people_links(job["company"])
    job["contacted"] = bool(c.execute(
        "SELECT 1 FROM outreach WHERE user=? AND lower(empresa)=lower(?)",
        (u, job["company"])).fetchone())
    n = company_postings(u).get((job["company"] or "").lower(), 1)
    job["company_postings"] = n
    job["size_hint"] = "xs" if n <= 2 else ("sm" if n <= 6 else "lg")
    job["favorite"] = 1 if job.get("favorite") == 1 else 0
    job["archived"] = 1 if job.get("archived") == 1 else 0
    job["salary"] = detect_salary(job.get("description") or "")
    job["tag_list"] = [t for t in (job.get("tags") or "").split(",") if t.strip()]
    gc = c.execute("SELECT after,created_at FROM gen_cvs WHERE user=? AND job_id=?",
                   (u, jid)).fetchone()
    job["cv_generated"] = ({"after": gc["after"], "created_at": gc["created_at"],
                            "download": f"/api/jobs/{jid}/cv/download"} if gc else None)
    return jsonify(job)


@app.route("/api/jobs/<jid>/messages")
@login_required
def api_job_messages(jid):
    u = cu()
    row = db().execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify([]), 404
    matched = [m for m in (row["matched"] or "").split(", ") if m]
    lang = "en" if request.args.get("lang") == "en" else "es"
    kws = [k.strip() for k in (request.args.get("keywords", "") or "").split(",") if k.strip()]
    return jsonify(recruiter_messages(user_profile(u), row["company"], row["title"],
                                      matched, top_cv_skills(u),
                                      request.args.get("contact", ""),
                                      lang=lang, keywords=kws))


@app.route("/api/jobs/<jid>/ai_message", methods=["POST"])
@login_required
def api_ai_message(jid):
    """Redacta un mensaje al reclutador desde cero con la IA configurada."""
    u = cu()
    row = db().execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    if get_config(u, "ai_enabled") != "1":
        return jsonify({"error": "La IA no está activada. Actívala en Ajustes → Redacción con IA."}), 400
    d = request.json or {}
    lang = "en" if d.get("lang") == "en" else "es"
    tone = d.get("tone") or "formal"
    kws = [k.strip() for k in (d.get("keywords", "") or "").split(",") if k.strip()]
    contact = (d.get("contact") or "").strip()
    cfg = {k: get_config(u, k) for k in ("ai_provider", "ai_model", "ai_api_key", "ai_base_url")}
    job = dict(row)
    want_variants = bool(d.get("variants"))
    if want_variants:
        system, user = ai_message_prompt(user_profile(u), job, lang, kws, "formal", contact)
        labels = ([("breve", "Breve (conexión)"), ("formal", "Formal"), ("cercano", "Cercano")])
        if lang == "en":
            system += ("\n\nProduce THREE distinct versions of the message: a SHORT connection note, "
                       "a FORMAL message, and a WARM value-first message. Separate them EXACTLY with a line "
                       "containing only '===NEXT==='. Output only the three messages and the separators.")
        else:
            system += ("\n\nProduce TRES versiones distintas del mensaje: una NOTA DE CONEXIÓN breve, "
                       "un mensaje FORMAL y un mensaje CERCANO con propuesta de valor. Sepáralas EXACTAMENTE con "
                       "una línea que contenga solo '===NEXT==='. Devuelve solo los tres mensajes y los separadores.")
    else:
        system, user = ai_message_prompt(user_profile(u), job, lang, kws, tone, contact)
    try:
        text = ai_chat(cfg, system, user, max_tokens=1200 if want_variants else 700)
        if not text:
            return jsonify({"error": "La IA no devolvió texto. Revisa el modelo configurado."}), 502
        if want_variants:
            parts = [p.strip() for p in re.split(r"={2,}\s*NEXT\s*={2,}", text) if p.strip()]
            titles = ["✨ Breve (conexión)", "✨ Formal", "✨ Cercano"]
            variants = [{"title": titles[i] if i < len(titles) else "✨ Variante",
                         "body": p} for i, p in enumerate(parts[:3])]
            if not variants:
                variants = [{"title": "✨ Redactado por IA", "body": text}]
            return jsonify({"ok": True, "variants": variants, "provider": cfg["ai_provider"]})
        return jsonify({"ok": True, "body": text, "provider": cfg["ai_provider"]})
    except requests.exceptions.ConnectionError:
        return jsonify({"error": "No se pudo conectar con la IA. Si usas Ollama, "
                        "verifica que esté corriendo (ollama serve) y el modelo descargado."}), 502
    except requests.exceptions.HTTPError as e:
        return jsonify({"error": f"La IA respondió con error: {e}. Revisa la API key y el modelo."}), 502
    except Exception as e:
        return jsonify({"error": f"Error de IA: {e}"}), 502


@app.route("/api/ai/test", methods=["POST"])
@login_required
def api_ai_test():
    """Prueba la conexión con la IA configurada (guarda antes lo que venga en el body)."""
    u = cu()
    d = request.json or {}
    for k in ("ai_enabled", "ai_provider", "ai_model", "ai_base_url"):
        if k in d:
            set_config(u, k, d[k])
    if d.get("ai_api_key"):
        set_config(u, "ai_api_key", d["ai_api_key"])
    cfg = {k: get_config(u, k) for k in ("ai_provider", "ai_model", "ai_api_key", "ai_base_url")}
    try:
        text = ai_chat(cfg, "Responde solo con la palabra: OK",
                       "Di OK", max_tokens=10)
        return jsonify({"ok": True, "reply": text[:80] or "(vacío)"})
    except Exception as e:
        return jsonify({"error": f"{e}"}), 502


# ============================================================================
#  GENERADOR DE KEYWORDS DE BÚSQUEDA (IA + respaldo)
# ============================================================================
def _ai_cfg(u):
    return {k: get_config(u, k) for k in ("ai_provider", "ai_model", "ai_api_key", "ai_base_url")}


def ai_json(cfg, system, user, max_tokens=600):
    """Llama a la IA y extrae el primer array/objeto JSON de la respuesta."""
    txt = ai_chat(cfg, system, user, max_tokens=max_tokens)
    m = re.search(r"\[.*\]|\{.*\}", txt, re.S)
    return json.loads(m.group(0) if m else txt)


# Enfoques base por área (respaldo sin IA) -----------------------------------
FOCUS_FALLBACK = {
    "ciberseguridad": ["Analista SOC / Blue Team", "Pentesting / Red Team",
                       "GRC / Cumplimiento", "Respuesta a incidentes / DFIR", "Seguridad en la nube"],
    "soporte": ["Soporte técnico N1", "Soporte técnico N2", "Administrador de sistemas",
                "Help Desk / Mesa de ayuda", "Soporte de redes"],
    "desarrollo": ["Frontend", "Backend", "Full Stack", "Móvil", "DevOps"],
    "datos": ["Analista de datos", "Ingeniero de datos", "Científico de datos", "BI / Reporting"],
    "redes": ["Administrador de redes", "Ingeniero de redes", "NOC", "Ciberseguridad de redes"],
}


def fallback_focuses(base):
    b = (base or "").lower()
    for key, vals in FOCUS_FALLBACK.items():
        if key in b:
            return vals
    base = base.strip() or "el área"
    return [base, f"{base} junior", f"{base} senior", f"{base} remoto"]


def _seniority(term, lang="es"):
    if lang == "en":
        return [term, f"Junior {term}", f"Senior {term}", f"{term} remote"]
    return [term, f"{term} junior", f"{term} senior", f"{term} remoto"]


# --- Scraping de keywords desde vacantes reales -----------------------------
KW_STOP = set(("de la el los las y o u en para con por un una que del al a e "
               "the and or of in for to with at por para en remoto remote hybrid "
               "híbrido presencial teletrabajo").split())


def _norm_title(t):
    t = re.sub(r"\(.*?\)", " ", t or "")
    t = re.sub(r"[|/•·¡!¿?,;:]+", " ", t)
    t = re.sub(r"[-–—]+", " ", t)
    t = re.sub(r"\b(jr|sr|ssr|junior|senior|semi\s*senior|trainee|becario|prácticas|remoto|remote|híbrido|hybrid)\b",
               "", t, flags=re.I)
    return re.sub(r"\s+", " ", t).strip()


def scrape_title_keywords(base, scopes, per=12):
    """Busca vacantes reales del término base en cada objetivo y extrae los
    títulos/expresiones más frecuentes como keywords que realmente dan con trabajos."""
    from collections import Counter
    norm_count, rep, countryof = Counter(), {}, {}
    bigrams, bigram_country = Counter(), {}
    got_any = False
    for label in scopes:
        loc = "" if label == "Remoto" else label
        try:
            cards = LI.search(base, loc, max_results=per)
        except Exception:
            cards = []
        for c in cards:
            got_any = True
            n = _norm_title(c.get("title") or "")
            key = n.lower()
            if len(key) < 4:
                continue
            norm_count[key] += 1
            rep.setdefault(key, n)
            countryof.setdefault(key, label)
            words = [w for w in re.findall(r"[a-záéíóúñ0-9\+\#\.]+", n.lower())
                     if w not in KW_STOP and len(w) > 2]
            for i in range(len(words) - 1):
                bg = words[i] + " " + words[i + 1]
                bigrams[bg] += 1
                bigram_country.setdefault(bg, label)
    if not got_any:
        return None   # sin red / sin resultados → deja usar el respaldo
    out, used = [], set()
    for key, _ in norm_count.most_common(14):
        if key not in used:
            used.add(key)
            out.append({"keyword": rep[key], "country": countryof[key]})
    for bg, cnt in bigrams.most_common(20):
        if cnt >= 2 and bg not in used and len(out) < 22:
            used.add(bg)
            out.append({"keyword": bg.title(), "country": bigram_country[bg]})
    return out[:22]


@app.route("/api/keywords/focuses", methods=["POST"])
@login_required
def api_keyword_focuses():
    """Sugiere enfoques/especializaciones concretos para un área."""
    u = cu()
    base = (request.json or {}).get("base", "").strip()
    if not base:
        return jsonify({"error": "Escribe qué buscas (ej. ciberseguridad)."}), 400
    if get_config(u, "ai_enabled") == "1":
        try:
            data = ai_json(_ai_cfg(u),
                           "Eres un experto en búsqueda de empleo. Devuelve SOLO un array JSON de strings, sin texto extra.",
                           f"El usuario busca trabajo en el área: '{base}'. Lista 5 enfoques o "
                           f"especializaciones concretas y buscables dentro de esa área, como etiquetas "
                           f"cortas (2-4 palabras). Devuelve solo el array JSON de strings.")
            focuses = [str(x).strip() for x in data if str(x).strip()][:6]
            if focuses:
                return jsonify({"focuses": focuses, "ai": True})
        except Exception:
            pass
    return jsonify({"focuses": fallback_focuses(base), "ai": False})


@app.route("/api/keywords/generate", methods=["POST"])
@login_required
def api_keyword_generate():
    """Genera keywords de búsqueda que dan con trabajos, localizadas por país/remoto."""
    u = cu()
    d = request.json or {}
    base = (d.get("base") or "").strip()
    focus = (d.get("focus") or base).strip()
    countries = [c.strip() for c in (d.get("countries") or []) if str(c).strip()]
    remote = bool(d.get("remote"))
    if not focus:
        return jsonify({"error": "Falta el enfoque o el término base."}), 400
    scopes = list(countries)
    if remote:
        scopes.append("Remoto")
    if not scopes:
        scopes = ["Remoto"]

    # 1) Scrapea títulos de vacantes reales del enfoque (la mejor señal)
    scraped = scrape_title_keywords(focus, scopes)

    # 2) Con IA: usa lo scrapeado como ejemplos reales para que localice y amplíe
    if get_config(u, "ai_enabled") == "1":
        examples = ""
        if scraped:
            examples = ("\n\nEjemplos REALES de títulos que aparecen ahora mismo en ofertas de esa búsqueda: "
                        + "; ".join(s["keyword"] for s in scraped[:12]) + ". Básate en ellos.")
        try:
            data = ai_json(_ai_cfg(u),
                           "Eres un experto en búsqueda de empleo internacional. Conoces cómo se nombran "
                           "los puestos en los portales de empleo de cada país. Devuelve SOLO JSON, sin texto extra.",
                           f"Genera términos de búsqueda (títulos de puesto y sinónimos que aparecen en ofertas "
                           f"reales) para encontrar trabajos de '{focus}' (área: {base}). "
                           f"Objetivos: {', '.join(scopes)}. Para cada objetivo usa el idioma y los términos "
                           f"locales de ese país (para 'Remoto' usa términos en inglés y español). "
                           f"Devuelve un array JSON de objetos {{\"keyword\":\"...\",\"country\":\"...\"}}, "
                           f"3-5 keywords por objetivo, sin duplicados.{examples}", max_tokens=900)
            out = []
            for it in data:
                kw = str(it.get("keyword", "")).strip()
                co = str(it.get("country", "")).strip() or scopes[0]
                if kw:
                    out.append({"keyword": kw, "country": co})
            if out:
                return jsonify({"suggestions": out[:40], "ai": True, "scraped": bool(scraped)})
        except Exception:
            pass

    # 3) Sin IA pero con scraping: devuelve las keywords reales extraídas
    if scraped:
        return jsonify({"suggestions": scraped, "ai": False, "scraped": True})

    # 4) Último respaldo (sin red / sin resultados): combos por seniority
    out = []
    for sc in scopes:
        lang = "en" if sc == "Remoto" else "es"
        for kw in _seniority(focus, lang):
            out.append({"keyword": kw, "country": sc})
    return jsonify({"suggestions": out[:40], "ai": False, "scraped": False})


@app.route("/api/board", methods=["GET", "POST"])
@login_required
def api_board():
    """Bloc de notas + pizarra guardados en la cuenta."""
    u = cu()
    if request.method == "POST":
        d = request.json or {}
        if "text" in d:
            set_config(u, "scratchpad", (d.get("text") or "")[:100000])
        if "board" in d:
            png = d.get("board") or ""
            if len(png) > 3_000_000:
                return jsonify({"error": "El dibujo es demasiado grande."}), 400
            set_config(u, "board_png", png)
        return jsonify({"ok": True})
    return jsonify({"text": get_config(u, "scratchpad"),
                    "board": get_config(u, "board_png")})


@app.route("/api/jobs/<jid>/cv", methods=["POST"])
@login_required
def api_job_cv(jid):
    import docx
    u = cu()
    row = db().execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    cvs = all_user_cvs(u)
    if not cvs:
        return jsonify({"error": "Primero carga tu CV en «Mi CV»."}), 400
    # Elegir la variante: la pedida, o la que mejor puntúa esta vacante
    label = (request.json or {}).get("cv") or row["best_cv"]
    cvsel = cv_by_label(u, label) if label else cvs[0]
    cv = cvsel["text"]
    ext = cvsel["ext"]
    orig = cvsel["orig"]
    extra = (request.json or {}).get("keywords", [])
    matched = [m for m in (row["matched"] or "").split(", ") if m]
    all_kw = list(dict.fromkeys(matched + extra))
    sd = user_skills(u)
    job_text = f"{row['title']} {row['description'] or ''}"
    # Score del CV ORIGINAL vs la vacante
    before = score_cv(profile_scoring_text(u, cv), job_text, sd)

    safe = re.sub(r"[^\w\-]+", "_", row["company"])[:30]
    out = os.path.join(GEN_DIR, f"CV_{u}_{safe}_{jid}.docx")

    if ext == ".docx" and os.path.exists(orig):
        info = build_ats_cv_preserve(orig, all_kw, out)
        preserved = True
    else:
        build_ats_cv(user_profile(u), cv, dict(row), matched, extra, out)
        info = {"added": extra, "where": "documento nuevo", "preserved": False}
        preserved = False

    # Score del CV GENERADO vs la vacante
    gen_text = _docx_all_text(docx.Document(out))
    after = score_cv(profile_scoring_text(u, gen_text), job_text, sd)

    # Guardar en el historial de CVs generados
    db().execute("""INSERT INTO gen_cvs(user,job_id,title,company,path,before,after,created_at)
        VALUES(?,?,?,?,?,?,?,?) ON CONFLICT(user,job_id) DO UPDATE SET
        path=excluded.path, before=excluded.before, after=excluded.after,
        created_at=excluded.created_at""",
        (u, jid, row["title"], row["company"], out,
         before["score"] if before else None, after["score"] if after else None,
         datetime.now().isoformat()))
    db().commit()

    return jsonify({
        "ok": True, "preserved": preserved,
        "before": before["score"] if before else None,
        "after": after["score"] if after else None,
        "added": info.get("added", []),
        "where": info.get("where", ""),
        "note": info.get("note", ""),
        "cv_used": cvsel["label"],
        "download": f"/api/jobs/{jid}/cv/download",
    })


@app.route("/api/jobs/<jid>/cv/download")
@login_required
def api_job_cv_download(jid):
    u = cu()
    row = db().execute("SELECT company FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    safe = re.sub(r"[^\w\-]+", "_", row["company"])[:30]
    out = os.path.join(GEN_DIR, f"CV_{u}_{safe}_{jid}.docx")
    if not os.path.exists(out):
        return jsonify({"error": "Genera el CV primero."}), 404
    return send_file(out, as_attachment=True, download_name=f"CV_{safe}.docx")


@app.route("/api/cvs")
@login_required
def api_cvs():
    rows = db().execute("""SELECT g.*, (g.path IS NOT NULL) has FROM gen_cvs g
                        WHERE g.user=? ORDER BY g.created_at DESC""", (cu(),)).fetchall()
    out = []
    for r in rows:
        d = dict(r)
        d["exists"] = os.path.exists(r["path"]) if r["path"] else False
        d["download"] = f"/api/jobs/{r['job_id']}/cv/download"
        out.append(d)
    return jsonify(out)


@app.route("/api/email/test", methods=["POST"])
@login_required
def api_email_test():
    cfg = email_config(cu())
    try:
        send_email(cfg, "Correo de prueba — Buscador de Vacantes",
                   "¡Funciona! Recibirás un correo así cuando una vacante encaje "
                   "con tu perfil por encima de tu umbral.")
    except Exception as e:
        return jsonify({"error": f"No se pudo enviar: {e}"}), 400
    return jsonify({"ok": True})


@app.route("/api/telegram/test", methods=["POST"])
@login_required
def api_telegram_test():
    u = cu()
    try:
        send_telegram(get_config(u, "telegram_token"), get_config(u, "telegram_chat_id"),
                      "✅ ¡Funciona! Recibirás aquí las vacantes que encajen con tu perfil.")
    except Exception as e:
        return jsonify({"error": f"No se pudo enviar: {e}"}), 400
    return jsonify({"ok": True})


@app.route("/api/jobs/<jid>/coverletter")
@login_required
def api_coverletter(jid):
    u = cu()
    row = db().execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    matched = [m for m in (row["matched"] or "").split(", ") if m]
    letter = cover_letter(user_profile(u), row["company"], row["title"], matched,
                          top_cv_skills(u), desc=row["description"] or "")
    return jsonify({"letter": letter})


@app.route("/api/jobs/<jid>/interview")
@login_required
def api_interview(jid):
    u = cu()
    row = db().execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    matched = [m for m in (row["matched"] or "").split(", ") if m]
    missing = [m for m in (row["missing"] or "").split(", ") if m]
    return jsonify(interview_prep(row["title"], matched, missing))


@app.route("/api/jobs/<jid>/requirements")
@login_required
def api_requirements(jid):
    u = cu()
    row = db().execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    cvsel = cv_by_label(u, row["best_cv"]) if row["best_cv"] else (all_user_cvs(u) or [{"text": ""}])[0]
    cv_text = cvsel["text"] if cvsel else ""
    return jsonify(requirements_breakdown(f"{row['title']} {row['description'] or ''}",
                                          cv_text, user_skills(u)))


@app.route("/api/followups")
@login_required
def api_followups():
    u = cu()
    days = int(get_config(u, "followup_days") or 5)
    cutoff = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d %H:%M")
    rows = db().execute(
        "SELECT * FROM outreach WHERE user=? AND responded=0 AND fecha<=? ORDER BY fecha ASC",
        (u, cutoff)).fetchall()
    return jsonify({"days": days, "items": [dict(r) for r in rows]})


@app.route("/api/metrics")
@login_required
def api_metrics():
    u = cu()
    c = db()
    # Embudo por estado
    funnel = {}
    for st in ("queued", "applied", "interview", "offer", "rejected"):
        funnel[st] = c.execute("SELECT COUNT(*) n FROM jobs WHERE user=? AND app_status=?",
                               (u, st)).fetchone()["n"]
    # Respuesta por portal (contactos vinculados a vacante)
    by_source = {}
    rows = c.execute("""SELECT j.source src, COUNT(*) tot, SUM(o.responded) resp
                     FROM outreach o LEFT JOIN jobs j ON j.id=o.job_id AND j.user=o.user
                     WHERE o.user=? GROUP BY j.source""", (u,)).fetchall()
    for r in rows:
        src = r["src"] or "otro"
        tot = r["tot"] or 0
        resp = r["resp"] or 0
        by_source[src] = {"total": tot, "responded": resp,
                          "rate": round(resp / tot * 100) if tot else 0}
    # Global
    tot = c.execute("SELECT COUNT(*) n FROM outreach WHERE user=?", (u,)).fetchone()["n"]
    resp = c.execute("SELECT COUNT(*) n FROM outreach WHERE user=? AND responded=1", (u,)).fetchone()["n"]
    return jsonify({"funnel": funnel, "by_source": by_source,
                    "contacted": tot, "responded": resp,
                    "rate": round(resp / tot * 100) if tot else 0})


# ---- Perfiles múltiples de CV ----
@app.route("/api/cv/variants", methods=["GET", "POST"])
@login_required
def api_cv_variants():
    u = cu()
    c = db()
    if request.method == "POST":
        f = request.files.get("file")
        label = (request.form.get("label") or "Variante").strip()
        if not f:
            return jsonify({"error": "No se recibió archivo."}), 400
        data = f.read()
        try:
            text = extract_cv_text(f.filename, data)
        except Exception as e:
            return jsonify({"error": str(e)}), 400
        ext = os.path.splitext(f.filename)[1].lower()
        cur = c.execute("INSERT INTO cv_variants(user,label,text,ext,filename,created) "
                        "VALUES(?,?,?,?,?,?)",
                        (u, label, text, ext, f.filename, datetime.now().isoformat()))
        vid = cur.lastrowid
        c.commit()
        try:
            with open(os.path.join(ORIG_DIR, f"{u}__{vid}{ext}"), "wb") as of:
                of.write(data)
        except Exception:
            pass
        _rescore_user(u)
        return jsonify({"ok": True})
    rows = c.execute("SELECT id,label,ext,filename,created FROM cv_variants WHERE user=?",
                     (u,)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/cv/variants/<int:vid>", methods=["DELETE"])
@login_required
def api_cv_variant_delete(vid):
    u = cu()
    db().execute("DELETE FROM cv_variants WHERE id=? AND user=?", (vid, u))
    db().commit()
    _rescore_user(u)
    return jsonify({"ok": True})


def _rescore_user(user):
    c = db()
    sd = user_skills(user)
    for row in c.execute("SELECT id,title,description FROM jobs WHERE user=?", (user,)).fetchall():
        sc, bestcv = best_score(user, f"{row['title']} {row['description'] or ''}", sd)
        if sc:
            c.execute("UPDATE jobs SET score=?,verdict=?,matched=?,missing=?,best_cv=? "
                      "WHERE user=? AND id=?",
                      (sc["score"], sc["verdict"], ", ".join(sc["matched"]),
                       ", ".join(sc["missing"]), bestcv, user, row["id"]))
    c.commit()


# ---- Vaciar y exportar ----
@app.route("/api/jobs/clear", methods=["POST"])
@login_required
def api_jobs_clear():
    u = cu()
    c = db()
    c.execute("DELETE FROM jobs WHERE user=?", (u,))
    c.execute("DELETE FROM alerts WHERE user=?", (u,))
    if (request.json or {}).get("searches"):
        c.execute("DELETE FROM searches WHERE user=?", (u,))
    c.commit()
    return jsonify({"ok": True})


@app.route("/api/telegram/digest_test", methods=["POST"])
@login_required
def api_digest_test():
    u = cu()
    conn = db()
    text = build_digest_text(u, conn, get_config(u, "digest_freq") or "daily")
    chan = get_config(u, "digest_channel") or "email"
    sent = []
    if chan in ("email", "both"):
        try:
            send_email(email_config(u), "[Vacantes] Resumen (prueba)", text)
            sent.append("correo")
        except Exception as e:
            return jsonify({"error": f"Correo: {e}"}), 400
    if chan in ("telegram", "both"):
        try:
            send_telegram(get_config(u, "telegram_token"), get_config(u, "telegram_chat_id"), text)
            sent.append("Telegram")
        except Exception as e:
            return jsonify({"error": f"Telegram: {e}"}), 400
    return jsonify({"ok": True, "sent": sent, "preview": text})


@app.route("/api/jobs/<jid>/status", methods=["POST"])
@login_required
def api_job_status(jid):
    st = (request.json or {}).get("status", "new")
    if st not in ("new", "queued", "applied", "interview", "offer", "rejected"):
        return jsonify({"error": "estado inválido"}), 400
    c = db()
    c.execute("UPDATE jobs SET app_status=? WHERE user=? AND id=?", (st, cu(), jid))
    # Registra la fecha de postulación la primera vez que llega a 'aplicado' o más allá
    if st in ("applied", "interview", "offer"):
        c.execute("UPDATE jobs SET applied_at=? WHERE user=? AND id=? AND "
                  "(applied_at IS NULL OR applied_at='')",
                  (datetime.now().isoformat(), cu(), jid))
    c.commit()
    return jsonify({"ok": True})


@app.route("/api/pipeline")
@login_required
def api_pipeline():
    rows = db().execute(
        "SELECT id,title,company,location,url,score,source,app_status FROM jobs "
        "WHERE user=? AND app_status IN ('queued','applied','interview','offer','rejected') "
        "ORDER BY (score IS NULL), score DESC", (cu(),)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/pipeline/clear", methods=["POST"])
@login_required
def api_pipeline_clear():
    """Saca del tablero las vacantes de un estado (o de todos) volviéndolas a 'new'."""
    st = (request.json or {}).get("status")
    c = db()
    if st in ("queued", "applied", "interview", "offer", "rejected"):
        c.execute("UPDATE jobs SET app_status='new' WHERE user=? AND app_status=?", (cu(), st))
    else:
        c.execute("UPDATE jobs SET app_status='new' WHERE user=? AND app_status IN "
                  "('queued','applied','interview','offer','rejected')", (cu(),))
    c.commit()
    return jsonify({"ok": True})


@app.route("/api/jobs/<jid>/assist", methods=["POST"])
@login_required
def api_assist(jid):
    """Aplicación asistida: prepara CV + carta + mensaje y marca 'applied'.
    NO envía nada: tú abres la vacante (ya logueado en tu navegador) y confirmas."""
    import docx
    u = cu()
    c = db()
    row = c.execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    job = dict(row)
    sd = user_skills(u)
    matched = [m for m in (job["matched"] or "").split(", ") if m]
    res = {"url": job["url"], "company": job["company"], "title": job["title"], "cv": None}
    cvs = all_user_cvs(u)
    if cvs:
        cvsel = cv_by_label(u, job.get("best_cv")) if job.get("best_cv") else cvs[0]
        cv = cvsel["text"]; ext = cvsel["ext"]; orig = cvsel["orig"]
        job_text = f"{job['title']} {job.get('description','') or ''}"
        before = score_cv(profile_scoring_text(u, cv), job_text, sd)
        safe = re.sub(r"[^\w\-]+", "_", job["company"])[:30]
        out = os.path.join(GEN_DIR, f"CV_{u}_{safe}_{jid}.docx")
        if ext == ".docx" and os.path.exists(orig):
            build_ats_cv_preserve(orig, matched, out)
        else:
            build_ats_cv(user_profile(u), cv, job, matched, [], out)
        after = score_cv(profile_scoring_text(u, _docx_all_text(docx.Document(out))), job_text, sd)
        c.execute("""INSERT INTO gen_cvs(user,job_id,title,company,path,before,after,created_at)
            VALUES(?,?,?,?,?,?,?,?) ON CONFLICT(user,job_id) DO UPDATE SET
            path=excluded.path, after=excluded.after, created_at=excluded.created_at""",
            (u, jid, job["title"], job["company"], out,
             before["score"] if before else None, after["score"] if after else None,
             datetime.now().isoformat()))
        res["cv"] = {"after": after["score"] if after else None,
                     "download": f"/api/jobs/{jid}/cv/download"}
    res["cover_letter"] = cover_letter(user_profile(u), job["company"], job["title"],
                                       matched, top_cv_skills(u), desc=job.get("description", ""))
    res["message"] = recruiter_messages(user_profile(u), job["company"], job["title"],
                                        matched, top_cv_skills(u))[2]["body"]
    c.execute("UPDATE jobs SET app_status='applied' WHERE user=? AND id=?", (u, jid))
    c.execute("UPDATE jobs SET applied_at=? WHERE user=? AND id=? AND (applied_at IS NULL OR applied_at='')",
              (datetime.now().isoformat(), u, jid))
    c.commit()
    return jsonify(res)


@app.route("/api/jobs/<jid>/autoapply", methods=["POST"])
@login_required
def api_autoapply(jid):
    """Prepara todo para aplicar de un clic: CV optimizado (solo con las
    keywords que YA tienes) + mensaje + enlace de la vacante."""
    import docx
    u = cu()
    c = db()
    row = c.execute("SELECT * FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    job = dict(row)
    cv = get_config(u, "cv_text")
    sd = user_skills(u)
    job_text = f"{job['title']} {job.get('description', '') or ''}"
    matched = [m for m in (job["matched"] or "").split(", ") if m]

    result = {"url": job["url"], "easy_apply": job["easy_apply"],
              "company": job["company"], "title": job["title"], "cv": None}

    if cv.strip():
        before = score_cv(profile_scoring_text(u, cv), job_text, sd)
        safe = re.sub(r"[^\w\-]+", "_", job["company"])[:30]
        out = os.path.join(GEN_DIR, f"CV_{u}_{safe}_{jid}.docx")
        ext = get_config(u, "cv_ext")
        orig = os.path.join(ORIG_DIR, f"{u}{ext}")
        if ext == ".docx" and os.path.exists(orig):
            build_ats_cv_preserve(orig, matched, out)
            preserved = True
        else:
            build_ats_cv(user_profile(u), cv, job, matched, [], out)
            preserved = False
        after = score_cv(profile_scoring_text(u, _docx_all_text(docx.Document(out))), job_text, sd)
        c.execute("""INSERT INTO gen_cvs(user,job_id,title,company,path,before,after,created_at)
            VALUES(?,?,?,?,?,?,?,?) ON CONFLICT(user,job_id) DO UPDATE SET
            path=excluded.path, before=excluded.before, after=excluded.after,
            created_at=excluded.created_at""",
            (u, jid, job["title"], job["company"], out,
             before["score"] if before else None, after["score"] if after else None,
             datetime.now().isoformat()))
        c.commit()
        result["cv"] = {"before": before["score"] if before else None,
                        "after": after["score"] if after else None,
                        "preserved": preserved,
                        "download": f"/api/jobs/{jid}/cv/download"}

    msgs = recruiter_messages(user_profile(u), job["company"], job["title"],
                              matched, top_cv_skills(u))
    result["message"] = msgs[2]["body"]  # cercano con propuesta de valor
    return jsonify(result)


# ---- CV ----
@app.route("/api/cv", methods=["POST"])
@login_required
def api_cv_upload():
    u = cu()
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No se recibió archivo."}), 400
    data = f.read()
    try:
        text = extract_cv_text(f.filename, data)
    except Exception as e:
        return jsonify({"error": str(e)}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    # Guardar el archivo ORIGINAL (para poder editarlo conservando el formato)
    try:
        with open(os.path.join(ORIG_DIR, f"{u}{ext}"), "wb") as of:
            of.write(data)
    except Exception:
        pass
    set_config(u, "cv_text", text)
    set_config(u, "cv_filename", f.filename)
    set_config(u, "cv_ext", ext)
    # Recalcular scores existentes con el MEJOR CV disponible
    c = db()
    sd = user_skills(u)
    for row in c.execute("SELECT id,title,description FROM jobs WHERE user=?", (u,)).fetchall():
        sc, bestcv = best_score(u, f"{row['title']} {row['description'] or ''}", sd)
        if sc:
            c.execute("UPDATE jobs SET score=?,verdict=?,matched=?,missing=?,best_cv=? "
                      "WHERE user=? AND id=?",
                      (sc["score"], sc["verdict"], ", ".join(sc["matched"]),
                       ", ".join(sc["missing"]), bestcv, u, row["id"]))
    c.commit()
    skills = sorted(detect_skills(text, sd))
    return jsonify({"ok": True, "filename": f.filename,
                    "words": len(text.split()), "skills": skills})


# ---- Alertas ----
@app.route("/api/alerts")
@login_required
def api_alerts():
    rows = db().execute("""SELECT a.id,a.score,a.created_at,a.read,j.title,j.company,
                        j.location,j.url,j.id as job_id FROM alerts a
                        JOIN jobs j ON j.id=a.job_id AND j.user=a.user
                        WHERE a.user=? ORDER BY a.created_at DESC LIMIT 200""",
                        (cu(),)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/alerts/<int:aid>/read", methods=["POST"])
@login_required
def api_alert_read(aid):
    db().execute("UPDATE alerts SET read=1 WHERE id=? AND user=?", (aid, cu()))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/alerts/read_all", methods=["POST"])
@login_required
def api_alerts_read_all():
    db().execute("UPDATE alerts SET read=1 WHERE user=?", (cu(),))
    db().commit()
    return jsonify({"ok": True})


# ---- Historial diario (para el gráfico) ----
@app.route("/api/history")
@login_required
def api_history():
    u = cu()
    c = db()
    jobs = {r["d"]: r["n"] for r in c.execute(
        "SELECT substr(found_at,1,10) d, COUNT(*) n FROM jobs WHERE user=? GROUP BY d",
        (u,)).fetchall()}
    matches = {r["d"]: r["n"] for r in c.execute(
        "SELECT substr(created_at,1,10) d, COUNT(*) n FROM alerts WHERE user=? GROUP BY d",
        (u,)).fetchall()}
    contacted = {r["d"]: r["n"] for r in c.execute(
        "SELECT substr(fecha,1,10) d, COUNT(*) n FROM outreach WHERE user=? GROUP BY d",
        (u,)).fetchall()}
    out = []
    today = datetime.now().date()
    for i in range(13, -1, -1):
        day = (today - timedelta(days=i)).isoformat()
        out.append({"date": day, "jobs": jobs.get(day, 0),
                    "matches": matches.get(day, 0),
                    "contacted": contacted.get(day, 0)})
    return jsonify(out)


# ---- Contactados ----
@app.route("/api/outreach", methods=["GET", "POST"])
@login_required
def api_outreach():
    u = cu()
    c = db()
    if request.method == "POST":
        d = request.json
        c.execute("""INSERT INTO outreach(user,empresa,rol,contacto,url,job_id,notas,fecha)
                     VALUES(?,?,?,?,?,?,?,?)""",
                  (u, d.get("empresa", ""), d.get("rol", ""), d.get("contacto", ""),
                   d.get("url", ""), d.get("job_id", ""), d.get("notas", ""),
                   datetime.now().strftime("%Y-%m-%d %H:%M")))
        c.commit()
        return jsonify({"ok": True})
    rows = c.execute("SELECT * FROM outreach WHERE user=? ORDER BY id DESC", (u,)).fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/outreach/<int:oid>", methods=["DELETE", "PATCH"])
@login_required
def api_outreach_edit(oid):
    u = cu()
    c = db()
    if request.method == "DELETE":
        c.execute("DELETE FROM outreach WHERE id=? AND user=?", (oid, u))
    else:
        d = request.json or {}
        if "notas" in d:
            c.execute("UPDATE outreach SET notas=? WHERE id=? AND user=?", (d["notas"], oid, u))
        if "responded" in d:
            c.execute("UPDATE outreach SET responded=? WHERE id=? AND user=?",
                      (1 if d["responded"] else 0, oid, u))
    c.commit()
    return jsonify({"ok": True})


@app.route("/api/export/xlsx")
@login_required
def api_export_xlsx():
    import openpyxl
    u = cu()
    c = db()
    wb = openpyxl.Workbook()

    def sheet(title, headers, rows):
        ws = wb.create_sheet(title)
        ws.append(headers)
        for r in rows:
            ws.append(r)

    wb.remove(wb.active)
    jobs = c.execute(
        "SELECT score,title,company,location,source,easy_apply,app_status,date,url,matched,missing "
        "FROM jobs WHERE user=? ORDER BY (score IS NULL), score DESC", (u,)).fetchall()
    sheet("Vacantes",
          ["Score", "Vacante", "Empresa", "Ubicación", "Portal", "Easy Apply",
           "Estado", "Fecha", "URL", "Coinciden", "Faltan"],
          [[j["score"], j["title"], j["company"], j["location"], j["source"],
            "Sí" if j["easy_apply"] == 1 else "", j["app_status"], j["date"], j["url"],
            j["matched"], j["missing"]] for j in jobs])
    pipe = c.execute("SELECT title,company,source,app_status,score FROM jobs WHERE user=? "
                     "AND app_status IN ('queued','applied','interview','offer','rejected')",
                     (u,)).fetchall()
    sheet("Postulaciones", ["Vacante", "Empresa", "Portal", "Estado", "Score"],
          [[p["title"], p["company"], p["source"], p["app_status"], p["score"]] for p in pipe])
    outr = c.execute("SELECT fecha,empresa,rol,contacto,responded,notas,url FROM outreach WHERE user=?",
                     (u,)).fetchall()
    sheet("Contactados", ["Fecha", "Empresa", "Rol", "Contacto", "Respondió", "Notas", "URL"],
          [[o["fecha"], o["empresa"], o["rol"], o["contacto"],
            "Sí" if o["responded"] else "No", o["notas"], o["url"]] for o in outr])
    al = c.execute("""SELECT a.created_at,j.title,j.company,a.score FROM alerts a
                   JOIN jobs j ON j.id=a.job_id AND j.user=a.user WHERE a.user=?""", (u,)).fetchall()
    sheet("Alertas", ["Fecha", "Vacante", "Empresa", "Score"],
          [[a["created_at"][:16], a["title"], a["company"], a["score"]] for a in al])
    cvg = c.execute("SELECT created_at,title,company,after FROM gen_cvs WHERE user=?", (u,)).fetchall()
    sheet("CVs generados", ["Fecha", "Vacante", "Empresa", "Score CV"],
          [[g["created_at"][:16], g["title"], g["company"], g["after"]] for g in cvg])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, as_attachment=True, download_name="historial_vacantes.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route("/api/outreach/export")
@login_required
def api_outreach_export():
    rows = db().execute("SELECT fecha,empresa,rol,contacto,url,notas FROM outreach "
                        "WHERE user=? ORDER BY id DESC", (cu(),)).fetchall()
    out = io.StringIO()
    w = csv.writer(out)
    w.writerow(["Fecha", "Empresa", "Rol", "Contacto", "URL", "Notas"])
    for r in rows:
        w.writerow([r["fecha"], r["empresa"], r["rol"], r["contacto"], r["url"], r["notas"]])
    return Response(out.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment; filename=contactados.csv"})


# ---- Config y diccionario de skills ----
@app.route("/api/config", methods=["POST"])
@login_required
def api_config():
    u = cu()
    d = request.json or {}
    for k, v in d.items():
        if k in DEFAULTS and k != "cv_text":
            if k in ("smtp_pass", "telegram_token", "adzuna_app_key", "ai_api_key") and not str(v).strip():
                continue  # no sobrescribir un secreto con vacío
            set_config(u, k, v)
    if "scan_enabled" in d:
        SCANNER.trigger()
    return jsonify({"ok": True})


@app.route("/api/skills", methods=["GET", "POST"])
@login_required
def api_skills():
    u = cu()
    if request.method == "POST":
        raw = (request.json or {}).get("skill_dict", "")
        # Validar que sea JSON de dict->lista (o vacío para usar el base)
        if raw.strip():
            try:
                d = json.loads(raw)
                assert isinstance(d, dict)
            except Exception:
                return jsonify({"error": "El diccionario debe ser JSON válido "
                               "(objeto {\"habilidad\": [\"variante\", ...]})."}), 400
        set_config(u, "skill_dict", raw)
        return jsonify({"ok": True})
    raw = get_config(u, "skill_dict")
    effective = parse_skill_dict(raw)
    return jsonify({"raw": raw, "is_custom": bool(raw.strip()),
                    "base": BASE_SKILLS, "effective": effective})


@app.route("/api/skills/reset", methods=["POST"])
@login_required
def api_skills_reset():
    set_config(cu(), "skill_dict", "")
    return jsonify({"ok": True})


@app.route("/api/password", methods=["POST"])
@login_required
def api_password():
    d = request.json or {}
    ok, msg = update_password(cu(), d.get("old", ""), d.get("new", ""))
    if not ok:
        return jsonify({"error": msg}), 400
    return jsonify({"ok": True})


@app.route("/api/jobs/export")
@login_required
def api_jobs_export():
    rows = db().execute("""SELECT score,easy_apply,title,company,location,date,url,
                        verdict,matched,missing,source_query,found_at
                        FROM jobs WHERE user=? ORDER BY (score IS NULL), score DESC""",
                        (cu(),)).fetchall()
    out = io.StringIO()
    w = csv.writer(out)
    w.writerow(["Score", "Solicitud sencilla", "Vacante", "Empresa", "Ubicación",
                "Fecha", "URL", "Veredicto", "Coinciden", "Faltan", "Búsqueda", "Encontrada"])
    for r in rows:
        w.writerow([r["score"], "Sí" if r["easy_apply"] == 1 else ("No" if r["easy_apply"] == 0 else ""),
                    r["title"], r["company"], r["location"], r["date"], r["url"],
                    r["verdict"], r["matched"], r["missing"], r["source_query"],
                    (r["found_at"] or "")[:10]])
    return Response("﻿" + out.getvalue(), mimetype="text/csv; charset=utf-8",
                    headers={"Content-Disposition": "attachment; filename=vacantes.csv"})


@app.route("/api/summary")
@login_required
def api_summary():
    u = cu()
    c = db()
    since = (datetime.now() - timedelta(days=7)).isoformat()
    since_d = (datetime.now() - timedelta(days=7)).strftime("%Y-%m-%d")
    found = c.execute("SELECT COUNT(*) n FROM jobs WHERE user=? AND found_at>=?",
                      (u, since)).fetchone()["n"]
    matches = c.execute("SELECT COUNT(*) n FROM alerts WHERE user=? AND created_at>=?",
                        (u, since)).fetchone()["n"]
    contacted = c.execute("SELECT COUNT(*) n FROM outreach WHERE user=? AND fecha>=?",
                          (u, since_d)).fetchone()["n"]
    best = c.execute("""SELECT title,company,score FROM jobs WHERE user=? AND found_at>=?
                     AND score IS NOT NULL ORDER BY score DESC LIMIT 1""",
                     (u, since)).fetchone()
    return jsonify({"found": found, "matches": matches, "contacted": contacted,
                    "best": dict(best) if best else None})


# ---- Favoritas ----
@app.route("/api/jobs/<jid>/favorite", methods=["POST"])
@login_required
def api_job_favorite(jid):
    u = cu()
    c = db()
    row = c.execute("SELECT favorite FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    newval = 0 if row["favorite"] == 1 else 1
    c.execute("UPDATE jobs SET favorite=? WHERE user=? AND id=?", (newval, u, jid))
    c.commit()
    return jsonify({"ok": True, "favorite": newval})


# ---- Entrevista + etiquetas por vacante ----
@app.route("/api/jobs/<jid>/meta", methods=["POST"])
@login_required
def api_job_meta(jid):
    u = cu()
    c = db()
    if not c.execute("SELECT 1 FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone():
        return jsonify({"error": "no encontrada"}), 404
    d = request.json or {}
    if "interview_at" in d:
        val = (d.get("interview_at") or "").strip()
        c.execute("UPDATE jobs SET interview_at=? WHERE user=? AND id=?", (val, u, jid))
    if "tags" in d:
        # Normaliza: separa por comas, quita vacíos y duplicados conservando orden
        raw = d.get("tags") or ""
        seen, clean = set(), []
        for t in raw.split(","):
            t = t.strip()
            key = t.lower()
            if t and key not in seen:
                seen.add(key)
                clean.append(t)
        c.execute("UPDATE jobs SET tags=? WHERE user=? AND id=?",
                  (",".join(clean[:12]), u, jid))
    if "notes" in d:
        c.execute("UPDATE jobs SET notes=? WHERE user=? AND id=?",
                  ((d.get("notes") or "").strip()[:4000], u, jid))
    c.commit()
    return jsonify({"ok": True})


# ---- Archivar / ocultar vacantes ----
@app.route("/api/jobs/<jid>/archive", methods=["POST"])
@login_required
def api_job_archive(jid):
    u = cu()
    c = db()
    row = c.execute("SELECT archived FROM jobs WHERE user=? AND id=?", (u, jid)).fetchone()
    if not row:
        return jsonify({"error": "no encontrada"}), 404
    newval = 0 if row["archived"] == 1 else 1
    c.execute("UPDATE jobs SET archived=? WHERE user=? AND id=?", (newval, u, jid))
    c.commit()
    return jsonify({"ok": True, "archived": newval})


@app.route("/api/upcoming")
@login_required
def api_upcoming():
    """Próximas entrevistas programadas (interview_at futuro o de hoy)."""
    u = cu()
    today = datetime.now().strftime("%Y-%m-%dT00:00")
    rows = db().execute(
        "SELECT id,title,company,url,interview_at,tags FROM jobs "
        "WHERE user=? AND interview_at IS NOT NULL AND interview_at!='' "
        "AND interview_at>=? ORDER BY interview_at ASC LIMIT 30",
        (u, today)).fetchall()
    return jsonify([dict(r) for r in rows])


# ---- Insights: skills a aprender + tendencias ----
@app.route("/api/insights")
@login_required
def api_insights():
    u = cu()
    c = db()
    # Skills a aprender: agrega los "missing" de TODAS las vacantes puntuadas
    rows = c.execute("SELECT missing FROM jobs WHERE user=? AND missing IS NOT NULL "
                     "AND missing!=''", (u,)).fetchall()
    learn = {}
    for r in rows:
        for s in (r["missing"] or "").split(", "):
            s = s.strip()
            if s:
                learn[s] = learn.get(s, 0) + 1
    learn_top = sorted(learn.items(), key=lambda kv: (-kv[1], kv[0]))[:15]
    total_scored = len(rows)
    # Tendencias: skills más demandadas (matched+missing) en los últimos 14 días
    since = (datetime.now() - timedelta(days=14)).isoformat()
    trows = c.execute("SELECT matched,missing FROM jobs WHERE user=? AND found_at>=?",
                      (u, since)).fetchall()
    demand = {}
    for r in trows:
        seen = set()
        for field in (r["matched"], r["missing"]):
            for s in (field or "").split(", "):
                s = s.strip()
                if s and s not in seen:
                    seen.add(s)
                    demand[s] = demand.get(s, 0) + 1
    trend_top = sorted(demand.items(), key=lambda kv: (-kv[1], kv[0]))[:12]
    return jsonify({
        "learn": [{"skill": k, "count": v} for k, v in learn_top],
        "learn_total": total_scored,
        "trends": [{"skill": k, "count": v} for k, v in trend_top],
        "trends_window": len(trows),
    })


# ---- Temas personalizados (guardados en la cuenta) ----
@app.route("/api/themes", methods=["GET", "POST"])
@login_required
def api_themes():
    u = cu()
    if request.method == "POST":
        d = request.json or {}
        name = (d.get("name") or "").strip()
        colors = d.get("colors") or {}
        if not name:
            return jsonify({"error": "Pon un nombre al tema."}), 400
        if not isinstance(colors, dict) or not colors:
            return jsonify({"error": "Faltan los colores del tema."}), 400
        try:
            themes = json.loads(get_config(u, "custom_themes") or "[]")
            assert isinstance(themes, list)
        except Exception:
            themes = []
        key = d.get("key") or ("custom-" + re.sub(r"[^a-z0-9]+", "-", name.lower()).strip("-")
                               or "custom-" + str(int(time.time())))
        # Reemplaza si ya existe una con la misma clave, si no la añade
        themes = [t for t in themes if t.get("key") != key]
        themes.append({"key": key, "name": name, "colors": colors})
        set_config(u, "custom_themes", json.dumps(themes[:40], ensure_ascii=False))
        return jsonify({"ok": True, "key": key})
    try:
        themes = json.loads(get_config(u, "custom_themes") or "[]")
        assert isinstance(themes, list)
    except Exception:
        themes = []
    return jsonify(themes)


@app.route("/api/themes/<key>", methods=["DELETE"])
@login_required
def api_theme_delete(key):
    u = cu()
    try:
        themes = json.loads(get_config(u, "custom_themes") or "[]")
        assert isinstance(themes, list)
    except Exception:
        themes = []
    themes = [t for t in themes if t.get("key") != key]
    set_config(u, "custom_themes", json.dumps(themes, ensure_ascii=False))
    return jsonify({"ok": True})


# ============================================================================
#  PERFIL DE USUARIO (con foto) + ONBOARDING
# ============================================================================
@app.route("/api/profile", methods=["GET", "POST"])
@login_required
def api_profile():
    u = cu()
    c = db()
    if request.method == "POST":
        d = request.json or {}
        fields = {}
        for k in ("display_name", "headline", "bio", "location", "links"):
            if k in d:
                fields[k] = str(d.get(k) or "").strip()[:2000]
        # Color de la palomita de verificación: solo admin/moderador/owner
        if "check_color" in d and (is_admin(u) or is_moderator(u)):
            col = str(d.get("check_color") or "").strip()
            if col == "" or re.match(r"^#[0-9a-fA-F]{6}$", col):
                fields["check_color"] = col
        if fields:
            sets = ", ".join(f"{k}=?" for k in fields)
            c.execute(f"UPDATE accounts SET {sets} WHERE username=?", (*fields.values(), u))
            c.commit()
        return jsonify({"ok": True})
    return jsonify(account_public(u))


@app.route("/api/profile/ai_bio", methods=["POST"])
@login_required
def api_profile_ai_bio():
    """Genera una bio profesional a partir de las respuestas de onboarding (si hay IA)."""
    u = cu()
    if get_config(u, "ai_enabled") != "1":
        return jsonify({"error": "Activa la IA en Ajustes para autogenerar tu bio."}), 400
    d = request.json or {}
    ans = d.get("answers") or {}
    cfg = {k: get_config(u, k) for k in ("ai_provider", "ai_model", "ai_api_key", "ai_base_url")}
    prompt = ("Con estas respuestas escribe una bio profesional en primera persona, natural y concisa "
              "(máx 3 frases), para el perfil de una plataforma de empleo. Solo la bio, sin comillas.\n\n"
              + json.dumps(ans, ensure_ascii=False))
    try:
        text = ai_chat(cfg, "Eres un redactor de perfiles profesionales. Escribe en español.", prompt, max_tokens=250)
        return jsonify({"ok": True, "bio": text.strip()})
    except Exception as e:
        return jsonify({"error": f"Error de IA: {e}"}), 502


@app.route("/api/profile/photo", methods=["POST"])
@login_required
def api_profile_photo():
    u = cu()
    f = request.files.get("photo")
    if not f or not f.filename:
        return jsonify({"error": "No se recibió imagen."}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        return jsonify({"error": "Formato no soportado (usa PNG, JPG, WEBP o GIF)."}), 400
    data = f.read()
    if len(data) > 4_000_000:
        return jsonify({"error": "La imagen es muy grande (máx 4 MB)."}), 400
    # limpia otras extensiones previas
    for e in (".png", ".jpg", ".jpeg", ".webp", ".gif"):
        p = os.path.join(AVATAR_DIR, u + e)
        if os.path.exists(p):
            try:
                os.remove(p)
            except Exception:
                pass
    with open(os.path.join(AVATAR_DIR, u + ext), "wb") as out:
        out.write(data)
    db().execute("UPDATE accounts SET avatar_ext=? WHERE username=?", (ext, u))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/profile/<username>/photo")
def api_profile_photo_get(username):
    a = db().execute("SELECT avatar_ext FROM accounts WHERE username=?", (username,)).fetchone()
    if a and a["avatar_ext"]:
        p = os.path.join(AVATAR_DIR, username + a["avatar_ext"])
        if os.path.exists(p):
            return send_file(p)
    return ("", 404)


@app.route("/api/profile/<username>")
@login_required
def api_profile_public(username):
    if not db().execute("SELECT 1 FROM accounts WHERE username=?", (username,)).fetchone():
        return jsonify({"error": "no existe"}), 404
    return jsonify(account_public(username))


# ============================================================================
#  ADMINISTRACIÓN: verificar usuarios y asignar roles
# ============================================================================
def _require_admin():
    if not is_admin(cu()):
        return jsonify({"error": "Solo administradores."}), 403
    return None


@app.route("/api/admin/users")
@login_required
def api_admin_users():
    guard = _require_admin()
    if guard:
        return guard
    rows = db().execute("SELECT username,role,verified,display_name,headline,created_at "
                        "FROM accounts ORDER BY (role='owner') DESC,(role='admin') DESC,created_at DESC").fetchall()
    return jsonify([dict(r) | {"has_avatar": bool(get_account(r["username"]).get("avatar_ext"))} for r in rows])


@app.route("/api/admin/verify", methods=["POST"])
@login_required
def api_admin_verify():
    guard = _require_admin()
    if guard:
        return guard
    d = request.json or {}
    target = (d.get("username") or "").strip()
    val = 1 if d.get("verified") else 0
    if not target:
        return jsonify({"error": "Falta el usuario."}), 400
    db().execute("UPDATE accounts SET verified=? WHERE username=?", (val, target))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/admin/role", methods=["POST"])
@login_required
def api_admin_role():
    guard = _require_admin()
    if guard:
        return guard
    d = request.json or {}
    target = (d.get("username") or "").strip()
    role = (d.get("role") or "").strip()
    if role not in ASSIGNABLE_ROLES:
        return jsonify({"error": "Rol inválido."}), 400
    if target.lower() == OWNER_USERNAME.lower():
        return jsonify({"error": "No se puede cambiar el rol del dueño."}), 400
    # solo el owner puede crear/quitar admins
    tgt = get_account(target)
    if (role == "admin" or tgt.get("role") == "admin") and get_account(cu()).get("role") != "owner":
        return jsonify({"error": "Solo el dueño puede gestionar administradores."}), 403
    db().execute("UPDATE accounts SET role=? WHERE username=?", (role, target))
    db().commit()
    return jsonify({"ok": True})


# ============================================================================
#  CHAT con salas/canales, presencia, MP y amigos
# ============================================================================
ROOM_RE = re.compile(r"^[a-z0-9\-]{2,24}$")


def touch_presence(u):
    db().execute("INSERT INTO presence(username,last_seen) VALUES(?,?) "
                 "ON CONFLICT(username) DO UPDATE SET last_seen=excluded.last_seen",
                 (u, datetime.now().isoformat()))
    db().commit()


def online_users(minutes=2):
    cutoff = (datetime.now() - timedelta(minutes=minutes)).isoformat()
    rows = db().execute("SELECT username FROM presence WHERE last_seen>=?", (cutoff,)).fetchall()
    return [r["username"] for r in rows]


@app.route("/api/rooms", methods=["GET", "POST"])
@login_required
def api_rooms():
    u = cu()
    c = db()
    if request.method == "POST":
        if not is_verified(u):
            return jsonify({"error": "Debes estar verificado para crear salas."}), 403
        name = (request.json or {}).get("name", "").strip().lower().replace(" ", "-")
        name = re.sub(r"[^a-z0-9\-]", "", name)
        if not ROOM_RE.match(name):
            return jsonify({"error": "Nombre de sala inválido (2-24, letras/números/guiones)."}), 400
        try:
            c.execute("INSERT INTO rooms(name,created_by,created_at) VALUES(?,?,?)",
                      (name, u, datetime.now().isoformat()))
            c.commit()
        except sqlite3.IntegrityError:
            return jsonify({"error": "Esa sala ya existe."}), 400
        return jsonify({"ok": True, "name": name})
    rows = c.execute("SELECT id,name,created_by FROM rooms ORDER BY (name='general') DESC, name ASC").fetchall()
    return jsonify([dict(r) for r in rows])


@app.route("/api/rooms/<int:rid>", methods=["DELETE"])
@login_required
def api_room_delete(rid):
    u = cu()
    row = db().execute("SELECT name,created_by FROM rooms WHERE id=?", (rid,)).fetchone()
    if not row:
        return jsonify({"error": "no existe"}), 404
    if row["name"] == "general":
        return jsonify({"error": "No se puede eliminar #general."}), 400
    if row["created_by"] != u and not is_moderator(u):
        return jsonify({"error": "No autorizado."}), 403
    db().execute("DELETE FROM rooms WHERE id=?", (rid,))
    db().execute("UPDATE chat SET deleted=1 WHERE room_id=?", (rid,))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/chat")
@login_required
def api_chat_get():
    u = cu()
    touch_presence(u)
    room = request.args.get("room", type=int) or 1
    after = request.args.get("after", type=int) or 0
    rows = db().execute("SELECT id,username,body,created_at FROM chat "
                        "WHERE deleted=0 AND room_id=? AND id>? ORDER BY id ASC LIMIT 200",
                        (room, after)).fetchall()
    out = []
    for r in rows:
        pub = account_public(r["username"])
        out.append({"id": r["id"], "username": r["username"], "body": r["body"],
                    "created_at": r["created_at"], "role": pub["role"], "verified": pub["verified"],
                    "display_name": pub["display_name"], "has_avatar": pub["has_avatar"],
                    "check_color": pub["check_color"]})
    online = [account_public(x) for x in online_users()]
    return jsonify({"messages": out, "me": cu(), "is_admin": is_admin(cu()),
                    "is_moderator": is_moderator(cu()), "can_write": is_verified(cu()),
                    "online": online})


@app.route("/api/chat", methods=["POST"])
@login_required
def api_chat_post():
    u = cu()
    if not is_verified(u):
        return jsonify({"error": "Tu cuenta debe estar verificada por un administrador para escribir."}), 403
    d = request.json or {}
    body = (d.get("body") or "").strip()
    room = int(d.get("room") or 1)
    if not body:
        return jsonify({"error": "Mensaje vacío."}), 400
    db().execute("INSERT INTO chat(username,body,created_at,room_id) VALUES(?,?,?,?)",
                 (u, body[:2000], datetime.now().isoformat(), room))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/chat/<int:mid>/delete", methods=["POST"])
@login_required
def api_chat_delete(mid):
    u = cu()
    row = db().execute("SELECT username FROM chat WHERE id=?", (mid,)).fetchone()
    if not row:
        return jsonify({"error": "no existe"}), 404
    if row["username"] != u and not is_moderator(u):
        return jsonify({"error": "Solo puedes borrar tus mensajes (o ser moderador)."}), 403
    db().execute("UPDATE chat SET deleted=1 WHERE id=?", (mid,))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/presence", methods=["POST"])
@login_required
def api_presence():
    touch_presence(cu())
    return jsonify({"online": [account_public(x) for x in online_users()]})


# ---- Amigos ----
def friendship(a, b):
    row = db().execute("SELECT requester,addressee,status FROM friends WHERE "
                       "(requester=? AND addressee=?) OR (requester=? AND addressee=?)",
                       (a, b, b, a)).fetchone()
    return dict(row) if row else None


@app.route("/api/friends")
@login_required
def api_friends():
    u = cu()
    rows = db().execute("SELECT requester,addressee,status FROM friends WHERE requester=? OR addressee=?",
                        (u, u)).fetchall()
    friends, incoming, outgoing = [], [], []
    online = set(online_users())
    for r in rows:
        other = r["addressee"] if r["requester"] == u else r["requester"]
        pub = account_public(other)
        pub["online"] = other in online
        if r["status"] == "accepted":
            friends.append(pub)
        elif r["addressee"] == u:
            incoming.append(pub)
        else:
            outgoing.append(pub)
    return jsonify({"friends": friends, "incoming": incoming, "outgoing": outgoing})


@app.route("/api/friends/request", methods=["POST"])
@login_required
def api_friend_request():
    u = cu()
    target = (request.json or {}).get("username", "").strip()
    if not target or target == u:
        return jsonify({"error": "Usuario inválido."}), 400
    if not db().execute("SELECT 1 FROM accounts WHERE username=?", (target,)).fetchone():
        return jsonify({"error": "Ese usuario no existe."}), 404
    fr = friendship(u, target)
    if fr:
        if fr["status"] == "accepted":
            return jsonify({"error": "Ya son amigos."}), 400
        # si el otro ya me había enviado solicitud, la aceptamos
        if fr["addressee"] == u:
            db().execute("UPDATE friends SET status='accepted' WHERE requester=? AND addressee=?",
                         (target, u))
            db().commit()
            return jsonify({"ok": True, "accepted": True})
        return jsonify({"error": "Solicitud ya enviada."}), 400
    db().execute("INSERT INTO friends(requester,addressee,status,created_at) VALUES(?,?,'pending',?)",
                 (u, target, datetime.now().isoformat()))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/friends/respond", methods=["POST"])
@login_required
def api_friend_respond():
    u = cu()
    d = request.json or {}
    other = (d.get("username") or "").strip()
    accept = bool(d.get("accept"))
    row = db().execute("SELECT 1 FROM friends WHERE requester=? AND addressee=? AND status='pending'",
                       (other, u)).fetchone()
    if not row:
        return jsonify({"error": "No hay solicitud pendiente."}), 404
    if accept:
        db().execute("UPDATE friends SET status='accepted' WHERE requester=? AND addressee=?", (other, u))
    else:
        db().execute("DELETE FROM friends WHERE requester=? AND addressee=?", (other, u))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/friends/remove", methods=["POST"])
@login_required
def api_friend_remove():
    u = cu()
    other = (request.json or {}).get("username", "").strip()
    db().execute("DELETE FROM friends WHERE (requester=? AND addressee=?) OR (requester=? AND addressee=?)",
                 (u, other, other, u))
    db().commit()
    return jsonify({"ok": True})


# ---- Mensajes privados ----
@app.route("/api/pm/threads")
@login_required
def api_pm_threads():
    """Lista de conversaciones (último mensaje por interlocutor)."""
    u = cu()
    rows = db().execute("SELECT sender,recipient,body,created_at,read FROM pms "
                        "WHERE sender=? OR recipient=? ORDER BY id DESC", (u, u)).fetchall()
    seen, threads = set(), []
    online = set(online_users())
    for r in rows:
        other = r["recipient"] if r["sender"] == u else r["sender"]
        if other in seen:
            continue
        seen.add(other)
        pub = account_public(other)
        unread = db().execute("SELECT COUNT(*) n FROM pms WHERE sender=? AND recipient=? AND read=0",
                              (other, u)).fetchone()["n"]
        threads.append({**pub, "online": other in online, "last": r["body"],
                        "last_at": r["created_at"], "unread": unread})
    return jsonify({"threads": threads})


@app.route("/api/pm/<other>")
@login_required
def api_pm_get(other):
    u = cu()
    db().execute("UPDATE pms SET read=1 WHERE sender=? AND recipient=?", (other, u))
    db().commit()
    rows = db().execute("SELECT id,sender,recipient,body,created_at FROM pms "
                        "WHERE (sender=? AND recipient=?) OR (sender=? AND recipient=?) "
                        "ORDER BY id ASC LIMIT 300", (u, other, other, u)).fetchall()
    return jsonify({"messages": [dict(r) for r in rows], "me": u, "other": account_public(other)})


@app.route("/api/pm/<other>", methods=["POST"])
@login_required
def api_pm_send(other):
    u = cu()
    if not is_verified(u):
        return jsonify({"error": "Debes estar verificado para enviar mensajes."}), 403
    if not db().execute("SELECT 1 FROM accounts WHERE username=?", (other,)).fetchone():
        return jsonify({"error": "Ese usuario no existe."}), 404
    body = (request.json or {}).get("body", "").strip()
    if not body:
        return jsonify({"error": "Mensaje vacío."}), 400
    db().execute("INSERT INTO pms(sender,recipient,body,created_at) VALUES(?,?,?,?)",
                 (u, other, body[:2000], datetime.now().isoformat()))
    db().commit()
    return jsonify({"ok": True})


# ============================================================================
#  BOLSA DE EMPLEO (empresas verificadas publican; verificados se postulan)
# ============================================================================
@app.route("/api/listings", methods=["GET", "POST"])
@login_required
def api_listings():
    u = cu()
    c = db()
    if request.method == "POST":
        if not is_verified(u):
            return jsonify({"error": "Tu cuenta debe estar verificada para publicar."}), 403
        if get_account(u).get("role") not in ("company", "admin", "owner"):
            return jsonify({"error": "Solo cuentas de tipo empresa pueden publicar vacantes. "
                            "Pide a un administrador que te asigne el rol de empresa."}), 403
        d = request.json or {}
        title = (d.get("title") or "").strip()
        if not title:
            return jsonify({"error": "Falta el título de la vacante."}), 400
        c.execute("""INSERT INTO listings(company,title,description,location,remote,salary,skills,created_at,active)
                     VALUES(?,?,?,?,?,?,?,?,1)""",
                  (u, title, (d.get("description") or "").strip(), (d.get("location") or "").strip(),
                   1 if d.get("remote") else 0, (d.get("salary") or "").strip(),
                   (d.get("skills") or "").strip(), datetime.now().isoformat()))
        c.commit()
        return jsonify({"ok": True})
    rows = c.execute("SELECT * FROM listings WHERE active=1 ORDER BY id DESC LIMIT 200").fetchall()
    out = []
    mine_apps = {r["listing_id"] for r in c.execute(
        "SELECT listing_id FROM applications WHERE applicant=?", (u,)).fetchall()}
    for r in rows:
        pub = account_public(r["company"])
        napp = c.execute("SELECT COUNT(*) n FROM applications WHERE listing_id=?", (r["id"],)).fetchone()["n"]
        out.append(dict(r) | {"company_name": pub["display_name"], "company_verified": pub["verified"],
                              "applicants": napp, "applied": r["id"] in mine_apps,
                              "is_owner": (r["company"] == u or is_admin(u))})
    return jsonify(out)


@app.route("/api/listings/<int:lid>", methods=["DELETE"])
@login_required
def api_listing_delete(lid):
    u = cu()
    row = db().execute("SELECT company FROM listings WHERE id=?", (lid,)).fetchone()
    if not row:
        return jsonify({"error": "no existe"}), 404
    if row["company"] != u and not is_admin(u):
        return jsonify({"error": "No autorizado."}), 403
    db().execute("UPDATE listings SET active=0 WHERE id=?", (lid,))
    db().commit()
    return jsonify({"ok": True})


@app.route("/api/listings/<int:lid>/apply", methods=["POST"])
@login_required
def api_listing_apply(lid):
    u = cu()
    if not is_verified(u):
        return jsonify({"error": "Tu cuenta debe estar verificada para postularte."}), 403
    row = db().execute("SELECT company FROM listings WHERE id=? AND active=1", (lid,)).fetchone()
    if not row:
        return jsonify({"error": "La vacante no existe."}), 404
    if row["company"] == u:
        return jsonify({"error": "No puedes postularte a tu propia vacante."}), 400
    msg = (request.json or {}).get("message", "").strip()[:2000]
    try:
        db().execute("INSERT INTO applications(listing_id,applicant,message,created_at) VALUES(?,?,?,?)",
                     (lid, u, msg, datetime.now().isoformat()))
        db().commit()
    except sqlite3.IntegrityError:
        return jsonify({"error": "Ya te postulaste a esta vacante."}), 400
    return jsonify({"ok": True})


@app.route("/api/listings/<int:lid>/applicants")
@login_required
def api_listing_applicants(lid):
    u = cu()
    row = db().execute("SELECT company FROM listings WHERE id=?", (lid,)).fetchone()
    if not row:
        return jsonify({"error": "no existe"}), 404
    if row["company"] != u and not is_admin(u):
        return jsonify({"error": "Solo la empresa que publicó puede ver los postulantes."}), 403
    apps = db().execute("SELECT applicant,message,created_at FROM applications WHERE listing_id=? ORDER BY id DESC",
                        (lid,)).fetchall()
    out = []
    for a in apps:
        pub = account_public(a["applicant"])
        cv = bool(get_config(a["applicant"], "cv_text"))
        out.append({"applicant": a["applicant"], "message": a["message"], "created_at": a["created_at"],
                    "display_name": pub["display_name"], "headline": pub["headline"], "bio": pub["bio"],
                    "location": pub["location"], "has_avatar": pub["has_avatar"],
                    "cv": cv, "cv_download": f"/api/applicants/{a['applicant']}/cv/download" if cv else None})
    return jsonify({"applicants": out})


@app.route("/api/applicants/<username>/cv/download")
@login_required
def api_applicant_cv(username):
    """La empresa que recibió una postulación (o un admin) puede ver el CV del postulante."""
    u = cu()
    authorized = is_admin(u) or db().execute(
        "SELECT 1 FROM applications a JOIN listings l ON a.listing_id=l.id "
        "WHERE a.applicant=? AND l.company=?", (username, u)).fetchone()
    if not authorized:
        return jsonify({"error": "No autorizado."}), 403
    # devuelve el CV original del postulante si existe
    ext = get_config(username, "cv_ext") or ".txt"
    for e in (ext, ".docx", ".pdf", ".txt"):
        p = os.path.join(ORIG_DIR, f"{username}{e}")
        if os.path.exists(p):
            return send_file(p, as_attachment=True, download_name=f"CV_{username}{e}")
    # respaldo: el texto plano del CV
    txt = get_config(username, "cv_text")
    if txt:
        return Response(txt, mimetype="text/plain",
                        headers={"Content-Disposition": f"attachment; filename=CV_{username}.txt"})
    return jsonify({"error": "El postulante no tiene CV cargado."}), 404


# ---- Favicon / estáticos personalizables ----
@app.route("/favicon.ico")
def favicon():
    for name in ("favicon.ico", "favicon.png", "favicon.svg"):
        p = os.path.join(STATIC_DIR, name)
        if os.path.exists(p):
            return send_file(p)
    return ("", 404)


@app.route("/static/<path:fname>")
def static_files(fname):
    p = os.path.join(STATIC_DIR, fname)
    if os.path.exists(p) and os.path.abspath(p).startswith(STATIC_DIR):
        return send_file(p)
    return ("", 404)


# ============================================================================
#  ARRANQUE
# ============================================================================

def open_browser():
    time.sleep(1.2)
    try:
        webbrowser.open("http://127.0.0.1:5000")
    except Exception:
        pass


if __name__ == "__main__":
    init_db()
    SCANNER.start()
    if os.environ.get("WERKZEUG_RUN_MAIN") != "true":
        threading.Thread(target=open_browser, daemon=True).start()
    # HOST=0.0.0.0 permite que otros equipos de tu red local entren (multiusuario).
    # Pon HOST=127.0.0.1 si quieres que sea solo tu máquina.
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", "5000"))
    print(f"\n  Buscador de Chamba — abre http://127.0.0.1:{port}")
    if host == "0.0.0.0":
        print("  En tu red local, otros entran por http://TU_IP_LOCAL:%d  (ej. http://192.168.1.50:%d)\n" % (port, port))
    app.run(host=host, port=port, debug=False)
